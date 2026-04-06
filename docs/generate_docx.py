"""
在庫管理システム ドキュメント Docx 生成スクリプト
python-docx を使用してMarkdownからWordファイルを生成
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_style(paragraph, level, text):
    """見出しスタイルを適用"""
    run = paragraph.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    else:
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)


def add_table(doc, table_lines):
    """Markdownテーブルをdocxテーブルに変換"""
    rows = []
    for line in table_lines:
        if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
            if ri == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                # ヘッダー背景色
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1E40AF')
                tcPr.append(shd)


def inline_text(text):
    """**bold** と `code` を除いたプレーンテキストを返す（docxでは個別にrun処理）"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def md_to_docx(md_text, doc):
    """MarkdownテキストをDocumentオブジェクトに追加"""
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # コードブロック
        if line.strip().startswith('```'):
            if in_code:
                if code_buf:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    run = p.add_run('\n'.join(code_buf))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)
                    code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 水平線
        if re.match(r'^---+\s*$', line) or re.match(r'^===+\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom = True
            i += 1
            continue

        # 見出し
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = inline_text(m.group(2))
            heading_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
            p = doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # テーブル
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            doc.add_paragraph()
            continue

        # 番号付きリスト
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m:
            text = inline_text(m.group(3))
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
            i += 1
            continue

        # 箇条書き
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            text = inline_text(m.group(2))
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            i += 1
            continue

        # 引用 (>)
        if line.strip().startswith('>'):
            text = inline_text(line.strip().lstrip('>').strip())
            p = doc.add_paragraph(style='Quote')
            p.add_run(text)
            i += 1
            continue

        # 通常段落
        para_lines = [line]
        i += 1
        while i < len(lines):
            nl = lines[i]
            if (not nl.strip() or nl.strip().startswith('#') or
                    nl.strip().startswith('|') or nl.strip().startswith('```') or
                    re.match(r'^(\s*)[-*]\s', nl) or re.match(r'^(\s*)\d+\.\s', nl) or
                    re.match(r'^---+\s*$', nl) or nl.strip().startswith('>')):
                break
            para_lines.append(nl)
            i += 1
        text = inline_text(' '.join(para_lines))
        doc.add_paragraph(text)

    if in_code and code_buf:
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_buf))
        run.font.name = 'Courier New'
        run.font.size = Pt(8)


def generate_docx(md_path, docx_path, doc_title):
    print(f"  生成中: {docx_path}")
    with open(md_path, encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()

    # ページ設定 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 既定フォントを日本語対応に設定
    doc.styles['Normal'].font.name = 'MS Gothic'
    doc.styles['Normal'].font.size = Pt(10)

    md_to_docx(md_text, doc)
    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) // 1024
    print(f"  完了: {docx_path} ({size_kb} KB)")


if __name__ == '__main__':
    BASE = os.path.dirname(os.path.abspath(__file__))
    docs = [
        ('要件定義書.md',  '要件定義書.docx'),
        ('仕様設計書.md',  '仕様設計書.docx'),
        ('取扱説明書.md',  '取扱説明書.docx'),
    ]

    for md_name, docx_name in docs:
        md_path   = os.path.join(BASE, md_name)
        docx_path = os.path.join(BASE, docx_name)
        generate_docx(md_path, docx_path, docx_name.replace('.docx', ''))

    print('\nDocx 3点 生成完了')
