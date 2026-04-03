# -*- coding: utf-8 -*-
"""
在庫管理システム 機能説明書（取扱説明書 + 仕様設計書 統合版）
サンプルデータを用いた実践的解説ドキュメント
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "在庫管理システム_機能説明書.docx"

# ─── XML ヘルパー ──────────────────────────────────────────────────────────
def _kwn(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))

def _kt(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepLines'))

def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit')
    cs.set(qn('w:val'), '1')
    trPr.append(cs)

def _row_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), '1')
    trPr.append(th)

def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _col_widths(table, widths_cm):
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])

# ─── 段落ヘルパー ──────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    _kwn(p); _kt(p)
    return p

def add_para(doc, text, kwn=False, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    _kt(p)
    if kwn: _kwn(p)
    return p

def add_note(doc, text, kwn=False):
    """薄グレーの注釈段落"""
    p = doc.add_paragraph()
    r = p.add_run(f"💡 {text}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0xAA)
    _kt(p)
    if kwn: _kwn(p)
    return p

def add_caution(doc, text):
    """赤系の注意書き"""
    p = doc.add_paragraph()
    r = p.add_run(f"⚠ {text}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    _kt(p)
    return p

def add_formula(doc, label, formula, example=None, kwn=False):
    """計算式ボックス"""
    p = doc.add_paragraph()
    _kt(p)
    if kwn: _kwn(p)
    r1 = p.add_run(f"【{label}】  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r2 = p.add_run(formula)
    r2.font.size = Pt(9)
    r2.font.name = 'Consolas'
    if example:
        r3 = p.add_run(f"   → 例: {example}")
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    return p

# ─── テーブルヘルパー ──────────────────────────────────────────────────────
def _make_table(doc, headers, rows, col_widths=None,
                header_color="1F4E79", alt_color="EBF3FB"):
    # anchor段落（keepNext付き）
    anc = doc.add_paragraph()
    anc.paragraph_format.space_before = Pt(0)
    anc.paragraph_format.space_after  = Pt(0)
    for child in list(anc._p):
        if child.tag != qn('w:pPr'):
            anc._p.remove(child)
    _kwn(anc)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダー行
    hdr = tbl.rows[0]
    _row_header(hdr); _cant_split(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        _shade(c, header_color)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # データ行
    for i, rd in enumerate(rows):
        row = tbl.rows[i + 1]
        _cant_split(row)
        bg = alt_color if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(rd):
            c = row.cells[j]
            _shade(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            txt = str(val)
            bold = txt.startswith("**") and txt.endswith("**")
            if bold: txt = txt[2:-2]
            r = c.paragraphs[0].add_run(txt)
            r.font.size = Pt(9)
            r.bold = bold

    if col_widths:
        _col_widths(tbl, col_widths)
    return tbl

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    return _make_table(doc, headers, rows, col_widths, header_color)

def add_sample_table(doc, rows3, intro=None, col_widths=None):
    """サンプルデータ付き3列テーブル（列名 | サンプル値 | 意味・説明）"""
    if intro:
        add_para(doc, intro, kwn=True)
    return _make_table(doc,
        ["列名 / 項目", "サンプル値", "説明"],
        rows3,
        col_widths or [4.2, 3.3, 9.0],
        header_color="2E75B6"
    )

def add_step_table(doc, steps, kwn_intro=False):
    """番号付きステップ表"""
    circles = ["①","②","③","④","⑤","⑥","⑦","⑧","⑨","⑩","⑪","⑫"]
    anc = doc.add_paragraph()
    anc.paragraph_format.space_before = Pt(0)
    anc.paragraph_format.space_after  = Pt(0)
    for child in list(anc._p):
        if child.tag != qn('w:pPr'):
            anc._p.remove(child)
    _kwn(anc)

    tbl = doc.add_table(rows=1 + len(steps), cols=3)
    tbl.style = 'Table Grid'
    headers = ["STEP", "処理名", "詳細"]
    hdr = tbl.rows[0]
    _row_header(hdr); _cant_split(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        _shade(c, "2E75B6")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    widths = [1.2, 3.5, 11.8]
    for i, (title, desc) in enumerate(steps):
        row = tbl.rows[i+1]
        _cant_split(row)
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for j, (val, w) in enumerate(zip([circles[i] if i < len(circles) else str(i+1), title, desc], widths)):
            c = row.cells[j]
            _shade(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.bold = True
            c.width = Cm(w)
    return tbl

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 表紙・目次
# ══════════════════════════════════════════════════════════════════════════════
def build_cover(doc):
    for _ in range(5): doc.add_paragraph()
    t = doc.add_paragraph("在庫管理システム")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(28)
    t.runs[0].bold = True
    t.runs[0].font.color.rgb = RGBColor(0x1F,0x4E,0x79)

    t2 = doc.add_paragraph("機能説明書")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(18)
    t2.runs[0].font.color.rgb = RGBColor(0x2E,0x75,0xB6)

    t3 = doc.add_paragraph("～ サンプルデータで学ぶ操作マニュアル＋仕様解説 ～")
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.runs[0].font.size = Pt(12)
    t3.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x66)
    t3.runs[0].italic = True

    for _ in range(3): doc.add_paragraph()
    t4 = doc.add_paragraph("Version 1.0　　2026年4月")
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t4.runs[0].font.size = Pt(11)
    t4.runs[0].font.color.rgb = RGBColor(0x44,0x44,0x44)
    page_break(doc)

def build_toc(doc):
    add_heading(doc, "目次", 1)
    items = [
        ("1","ログイン・認証"),
        ("2","ダッシュボード"),
        ("3","在庫管理"),
        ("　3.1","在庫一覧・検索"),
        ("　3.2","在庫編集"),
        ("　3.3","在庫廃棄・復元"),
        ("　3.4","ロケーション移動"),
        ("4","商品マスタ管理"),
        ("　4.1","商品一覧・CRUD"),
        ("　4.2","一括インポート / エクスポート"),
        ("5","発注管理"),
        ("　5.1","手動発注"),
        ("　5.2","🔄 自動チェック実行"),
        ("　5.3","混載グループ管理"),
        ("　5.4","発注残・部分入荷管理"),
        ("6","需要予測"),
        ("　6.1","商品別 予測一覧（画面の見方）"),
        ("　6.2","計算ロジック詳解"),
        ("　6.3","発注点・発注数の自動反映"),
        ("　6.4","販促予定・受注予定の登録"),
        ("7","チェーン・店舗マスタ"),
        ("8","CSV自動インポート"),
        ("9","メール通知"),
        ("10","権限管理"),
        ("付録","主要テーブル一覧"),
    ]
    for num, title in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0 if not num.startswith("　") else 1)
        r = p.add_run(f"{num}　{title}")
        r.font.size = Pt(10)
        r.bold = not num.startswith("　") and num != "付録"
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. ログイン・認証
# ══════════════════════════════════════════════════════════════════════════════
def ch01_login(doc):
    add_heading(doc, "1.　ログイン・認証", 1)

    add_heading(doc, "1.1　ログイン画面の操作", 2)
    add_para(doc, "ブラウザで http://localhost:5000 にアクセスするとログイン画面が表示されます。", kwn=True)
    add_sample_table(doc, [
        ["ユーザー名", "admin", "システム管理者から発行されたユーザー名"],
        ["パスワード", "●●●●●●", "初回ログイン時は管理者から通知された仮パスワード"],
    ])

    add_heading(doc, "1.2　初回ログイン時のパスワード変更", 2)
    add_para(doc, "初回ログインのみパスワード変更画面へ強制移動します。新しいパスワードを設定するまで他の画面は利用できません。", kwn=True)
    add_step_table(doc, [
        ("ログイン", "ユーザー名・仮パスワードを入力してログインボタンをクリック"),
        ("パスワード変更画面表示", "自動的にパスワード変更画面へリダイレクトされる"),
        ("新パスワード入力", "新しいパスワードを2回入力して「変更する」ボタンをクリック"),
        ("通常画面へ移動", "ダッシュボードへ自動移動。次回から新パスワードでログイン可能"),
    ])

    add_heading(doc, "1.3　セキュリティ仕様", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["項目", "仕様"],
        [
            ["レート制限", "同一IPから10回/分を超えるログイン試行は自動ブロック（429エラー）"],
            ["パスワード保存", "bcryptハッシュ化して保存。平文は一切保存しない"],
            ["セッション管理", "ログイン成功後にFlaskセッションを発行。ブラウザを閉じると無効化"],
            ["権限管理", "ユーザーごとに閲覧・操作できる画面を個別に設定可能"],
        ],
        col_widths=[4.5, 12.0]
    )
    add_caution(doc, "ログインに連続して失敗するとアクセスが一時的にロックされます。しばらく待ってから再試行してください。")
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ダッシュボード
# ══════════════════════════════════════════════════════════════════════════════
def ch02_dashboard(doc):
    add_heading(doc, "2.　ダッシュボード", 1)
    add_para(doc, "ログイン後に最初に表示されるトップページ。在庫・発注・賞味期限の状態をひと目で確認できます。")

    add_heading(doc, "2.1　サマリカード（画面上部）", 2)
    add_sample_table(doc, [
        ["有効商品数", "247", "is_active=1（有効）の商品数。無効化・削除した商品は含まない"],
        ["発注アラート", "8", "発注点以下かつ当日未発注の商品数。赤数字で表示"],
        ["賞味期限アラート", "3", "30日以内に賞味期限を迎えるロット数。黄数字で表示"],
    ])

    add_heading(doc, "2.2　発注アラートの判定条件", 2)
    add_para(doc, "以下の条件をすべて満たす商品がアラートとしてカウントされます。", kwn=True)
    add_table(doc,
        ["条件", "サンプル", "意味"],
        [
            ["is_active = 1", "有効", "有効商品のみ対象"],
            ["reorder_point > 0", "50 個", "発注点が設定されていること"],
            ["現在庫 ≤ 発注点", "在庫30 ≤ 発注点50", "在庫が発注点以下に落ちている"],
            ["ordered_at IS NULL", "本日未発注", "同日中にすでに発注済みの場合は除外"],
        ],
        col_widths=[4.5, 3.5, 8.5]
    )

    add_heading(doc, "2.3　直近発注履歴", 2)
    add_sample_table(doc, [
        ["発注日", "2026-04-03", "発注を行った日付"],
        ["商品名", "天日干したくあん", "発注商品"],
        ["数量", "20", "発注数量（個またはケース）"],
        ["メール結果", "送信済み", "仕入先へのメール送信結果。失敗時は「エラー」と表示"],
    ])
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. 在庫管理
# ══════════════════════════════════════════════════════════════════════════════
def ch03_inventory(doc):
    add_heading(doc, "3.　在庫管理", 1)

    # 3.1
    add_heading(doc, "3.1　在庫一覧・検索", 2)
    add_para(doc, "現在の在庫をロット・ロケーション単位で一覧表示します。検索ボックスにJAN・商品名などを入力して絞り込めます。", kwn=True)
    add_sample_table(doc, [
        ["JAN", "4971635544602", "JANコード。バーコードと同じ番号"],
        ["商品名", "天日干したくあん", "商品マスタに登録した商品名"],
        ["仕入先", "キムラ漬物", "仕入先名"],
        ["ロケーション", "A-01-03", "保管場所コード（棚番号など）"],
        ["数量", "48", "このロットの在庫数量（個）"],
        ["賞味期限", "2026-09-30", "このロットの賞味期限。近づくと赤表示"],
        ["ロット番号", "L2026031501", "入荷ロットを識別する番号"],
    ])
    add_note(doc, "在庫は「1ロット = 1行」で表示されます。同じ商品でも賞味期限・ロケーションが異なれば別行になります。")

    # 3.2
    add_heading(doc, "3.2　在庫直接編集", 2)
    add_para(doc, "在庫数量・賞味期限・ロケーションを直接変更します。変更内容はすべて変動履歴として記録されます。", kwn=True)
    add_step_table(doc, [
        ("在庫を選択", "在庫一覧で対象ロットの「編集」ボタンをクリック"),
        ("値を変更", "数量・賞味期限・ロケーションを変更。例：数量 48 → 45"),
        ("保存", "「保存」ボタンをクリック。変更前後の数量が stock_movements に記録される"),
    ])
    add_note(doc, "在庫を0以下に設定するとレコードが削除されます（在庫なし扱い）。廃棄の場合は「廃棄」機能を使用してください。")

    # 3.3
    add_heading(doc, "3.3　在庫廃棄・復元", 2)
    add_para(doc, "賞味期限切れ・破損などで販売できない在庫を廃棄します。誤廃棄の場合は復元できます。", kwn=True)
    add_sample_table(doc, [
        ["廃棄理由", "expiry", "expiry=賞味期限切れ / damage=破損 / other=その他"],
        ["廃棄数量", "12", "廃棄する個数"],
        ["廃棄ロス金額", "1,440円", "廃棄数量 × 原価（原価120円の場合: 12×120=1,440円）"],
    ])

    add_para(doc, "")
    add_formula(doc, "廃棄ロス金額",
        "廃棄数量 × 原価（cost_price）",
        "12個 × 120円 = 1,440円")

    add_table(doc,
        ["操作", "説明"],
        [
            ["廃棄実行", "stocks から該当ロットを削除し disposed_stocks へ移動。stock_movements に廃棄ログを記録"],
            ["復元", "disposed_stocks から stocks へ戻す。誤廃棄の修正に使用"],
        ],
        col_widths=[4.0, 12.5]
    )

    # 3.4
    add_heading(doc, "3.4　ロケーション移動", 2)
    add_para(doc, "在庫を別の保管場所へ移します。移動元・移動先の在庫が自動的に更新されます。", kwn=True)
    add_step_table(doc, [
        ("移動元を選択", "在庫一覧で移動したいロットの「移動」ボタンをクリック"),
        ("移動先と数量を入力", "移動先ロケーション（例: B-02-01）と移動数量（例: 20個）を入力"),
        ("数量チェック", "移動数量 > 移動元在庫数量の場合はエラー。20 > 48 でないので OK"),
        ("在庫更新", "移動元: 48 → 28個に減算。移動先に同一JANが存在すれば加算、なければ新規作成"),
        ("移動ログ記録", "stock_transfers テーブルに移動元・移動先・数量を記録"),
    ])
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. 商品マスタ管理
# ══════════════════════════════════════════════════════════════════════════════
def ch04_products(doc):
    add_heading(doc, "4.　商品マスタ管理", 1)

    add_heading(doc, "4.1　商品一覧・登録・編集", 2)
    add_para(doc, "商品マスタはシステム全体の基盤です。発注・予測・在庫管理に必要な全パラメータをここで設定します。", kwn=True)
    add_sample_table(doc, [
        ["JAN", "4971635544602", "JANコード（主キー）。バーコードの数字13桁"],
        ["商品CD", "1220005", "仕入先が管理する商品コード。発注書に記載"],
        ["商品名", "天日干したくあん", "商品名（画面・発注書・メールに使用）"],
        ["仕入先CD", "122", "仕入先識別コード"],
        ["仕入先名", "キムラ漬物", "仕入先会社名"],
        ["仕入先メール", "order@kimura-tsukemono.co.jp", "自動発注メールの送信先アドレス"],
        ["入数（unit_qty）", "6", "1ケースあたりの個数。発注はケース単位に切り上げ"],
        ["発注数量（order_qty）", "40", "標準の発注数量（個）"],
        ["発注点（reorder_point）", "30", "この数以下になったら発注トリガー"],
        ["発注モード（reorder_auto）", "1", "0=手動 / 1=AIモード / 2=前年実績モード"],
        ["リードタイム（lead_time_days）", "5", "発注から入荷までの日数"],
        ["安全係数（safety_factor）", "1.2", "予測需要に掛ける余裕係数。1.2 = 20%増し"],
        ["賞味期限日数（shelf_life_days）", "365", "入荷日からの賞味期限（日）"],
        ["アラート閾値（expiry_alert_days）", "30", "残り30日以内になるとアラート表示"],
        ["混載グループ（mixed_group）", "KIM-LOT1", "同グループ商品をまとめてロット発注する際に設定"],
        ["原価（cost_price）", "120", "廃棄ロス計算や発注金額の計算に使用"],
        ["有効フラグ（is_active）", "1", "1=有効・0=無効（論理削除）"],
    ])

    add_heading(doc, "4.2　発注モード（reorder_auto）の使い分け", 2)
    add_para(doc, "商品ごとに最適な発注点の管理方法を選べます。", kwn=True)
    add_table(doc,
        ["値", "モード", "こんなときに使う"],
        [
            ["0", "手動", "担当者が自分で発注点を決めたい場合。自動チェックの対象外"],
            ["1", "AIモード", "過去売上データが十分ある主力商品。曜日・季節を加味した自動更新"],
            ["2", "前年実績モード", "季節商品・催事品。前年同月の実績を基に発注点を自動設定"],
        ],
        col_widths=[1.5, 3.5, 11.5]
    )

    add_heading(doc, "4.3　一括インポート", 2)
    add_para(doc, "Excel・CSV形式で複数商品を一括登録・更新できます。", kwn=True)
    add_step_table(doc, [
        ("テンプレートDL", "「Excelテンプレート」ボタンからテンプレートをダウンロード"),
        ("データ入力", "テンプレートに商品情報を入力。発注モード列はドロップダウンで選択可"),
        ("インポート", "「ファイルを選択」→「インポート」ボタンをクリック"),
        ("モード選択", "upsert（既存更新＋新規追加）または add_only（新規追加のみ）を選択"),
        ("結果確認", "成功件数・スキップ件数がフラッシュメッセージで表示される"),
    ])
    add_note(doc, "JANコードが 4.9E+12 のような指数表記になっている場合も自動で正規化されます。")
    add_note(doc, "日付形式は YYYY/MM/DD・YYYYMMDD・YY/MM/DD など10種類以上を自動判定します。")
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. 発注管理
# ══════════════════════════════════════════════════════════════════════════════
def ch05_orders(doc):
    add_heading(doc, "5.　発注管理", 1)

    # 5.1
    add_heading(doc, "5.1　手動発注", 2)
    add_para(doc, "画面上で商品を選んで手動で発注します。発注書メールが仕入先へ自動送信されます。", kwn=True)
    add_step_table(doc, [
        ("発注チェック画面を開く", "メニューから「発注管理」→「発注チェック」をクリック"),
        ("商品を選択", "発注対象商品の一覧からチェックボックスを選択。例：天日干したくあん"),
        ("数量確認", "発注数量（デフォルト: order_qty = 40）を確認。必要に応じて変更"),
        ("ケース単位丸め", "unit_qty=6の場合、40÷6=6.67 → 7ケース（42個）に自動切り上げ"),
        ("発注実行", "「発注する」ボタンをクリック"),
        ("メール送信", "仕入先「キムラ漬物」のメールアドレスへ発注書を自動送信"),
        ("発注済みフラグ", "products.ordered_at に本日日付をセット。当日の重複発注を防止"),
    ])
    add_formula(doc, "ケース単位切り上げ",
        "ceil(発注数量 ÷ 入数) × 入数",
        "ceil(40 ÷ 6) × 6 = 7 × 6 = 42個")
    add_caution(doc, "当日すでに発注済みの商品は「発注済み」タブに移動します。同日中は再発注できません。")

    # 5.2
    add_heading(doc, "5.2　🔄 自動チェック実行", 2)
    add_para(doc, "ボタン操作または毎朝8時のスケジューラーが自動で発注要否を判断し、条件を満たした商品を自動発注します。", kwn=True)

    add_table(doc,
        ["トリガー方法", "タイミング"],
        [
            ["手動実行", "発注チェック画面の「🔄 自動チェック実行」ボタンをクリック"],
            ["スケジューラー（自動）", "毎日朝8:00にバックグラウンドで自動実行（設定画面で時刻変更可）"],
        ],
        col_widths=[5.0, 11.5]
    )

    add_para(doc, "")
    add_para(doc, "【処理ステップ】", bold=True, kwn=True)
    add_step_table(doc, [
        ("対象商品の抽出", "reorder_auto = 1（AI）または 2（前年実績）かつ is_active=1 の全商品をループ"),
        ("有効在庫の計算", "有効在庫 = 現在庫 − 賞味期限アラート以内の在庫（例: 48 − 0 = 48）"),
        ("重複チェック", "当日すでに order_history に記録あり → スキップ。order_pending に保留中 → スキップ"),
        ("発注トリガー判定", "有効在庫 ≤ 発注点（30）の場合: 発注点到達。在庫 ≥ メーカーロット数の場合: ロット数到達"),
        ("発注数量の計算", "不足数 = 発注点 − 有効在庫 = 30 − 10 = 20。ceil(20 ÷ order_qty) × order_qty"),
        ("混載グループ判定", "mixed_group 設定あり → order_pending に「保留」登録。なし → 即時発注"),
        ("混載条件チェック", "グループ合計が閾値以上 or 強制送信日を超過した場合に送信"),
        ("メール一括送信", "仕入先ごとに1通の発注書メールを送信"),
        ("DB記録更新", "order_history・order_pending・products.ordered_at・alert_logs を更新"),
    ])

    add_para(doc, "")
    add_para(doc, "【発注チェック画面の3タブ】", bold=True, kwn=True)
    add_table(doc,
        ["タブ名", "表示条件", "サンプル"],
        [
            ["発注対象商品", "在庫が発注点以下で今すぐ発注が必要", "天日干したくあん（在庫10 ≤ 発注点30）"],
            ["混載ロット保留中", "混載グループの合計ケース数が未達で保留中", "キムラ漬物グループ（あと2ケース不足）"],
            ["発注済み", "本日すでに ordered_at がセットされている", "千枚漬け（本日9:00発注済み）"],
        ],
        col_widths=[4.0, 5.5, 7.0]
    )

    # 5.3
    add_heading(doc, "5.3　混載グループ管理", 2)
    add_para(doc, "複数商品をまとめて1トラック分のロットで発注する「混載発注」を管理します。グループ合計が設定ケース数に達したら発注します。", kwn=True)
    add_sample_table(doc, [
        ["mixed_group", "KIM-LOT1", "グループ名。同名の商品が同一グループとして管理される"],
        ["mixed_lot_qty", "10", "発注条件となる合計ケース数（10ケース揃ったら発注）"],
        ["mixed_mode", "gte", "gte（10ケース以上）/ unit（10の倍数）"],
        ["mixed_force_days", "3", "保留開始から3日後には条件未達でも強制送信"],
        ["強制送信日", "2026-04-06", "保留開始(4/3) + 3日 = 4/6に強制送信"],
    ])

    # 5.4
    add_heading(doc, "5.4　発注残・部分入荷管理", 2)
    add_para(doc, "発注済みで未入荷の商品を管理します。部分入荷（分割納品）にも対応しています。", kwn=True)
    add_sample_table(doc, [
        ["発注日", "2026-04-01", "発注を行った日"],
        ["商品名", "天日干したくあん", "対象商品"],
        ["発注数量", "42", "発注した数量（個）"],
        ["入荷済み数量", "18", "すでに受領登録済みの数量"],
        ["発注残数量", "24", "42 − 18 = まだ届いていない数量"],
        ["予定入荷日", "2026-04-06", "expected_receipt_date。リードタイム5日から自動計算"],
        ["遅延区分", "延滞", "予定入荷日を過ぎると「延滞」（赤）、間近は「要注意」（橙）"],
    ])

    add_para(doc, "")
    add_para(doc, "【受領登録の手順】", bold=True, kwn=True)
    add_step_table(doc, [
        ("発注残一覧を開く", "メニューから「発注残管理」をクリック"),
        ("対象商品を選択", "「受領登録」ボタンをクリック"),
        ("受領情報を入力", "入荷数量（18）・賞味期限（2026-09-30）・ロット番号・ロケーションを入力"),
        ("登録実行", "stocks テーブルに在庫が追加され、stock_movements に入荷ログが記録される"),
        ("部分入荷の場合", "発注残（24個）が引き続き一覧に残る。残りが届いたら再度受領登録"),
        ("完了処理", "全量入荷またはキャンセル時に「クローズ」ボタンで発注残を終了"),
    ])
    add_formula(doc, "発注残数量", "発注数量 − 入荷済み数量", "42 − 18 = 24個")
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. 需要予測（メイン章）
# ══════════════════════════════════════════════════════════════════════════════
def ch06_forecast(doc):
    add_heading(doc, "6.　需要予測", 1)
    add_para(doc, "過去の売上データを分析し、今後30日の需要を予測します。予測結果をそのまま発注点・発注数の設定に反映できます。")

    # 6.1
    add_heading(doc, "6.1　商品別 予測一覧 ─ 画面の見方", 2)
    add_para(doc, "メニューから「需要予測」をクリックすると「商品別 予測一覧」が表示されます。以下のサンプル商品を例に各列を説明します。", kwn=True)

    # 商品情報ブロック
    add_para(doc, "【サンプル商品情報】", bold=True, kwn=True)
    add_table(doc,
        ["表示内容", "サンプル値", "意味"],
        [
            ["仕入先名 / 仕入先CD", "キムラ漬物 / 122", "仕入先を識別する名前とコード"],
            ["商品名", "天日干したくあん", "商品マスタに登録した商品名"],
            ["商品CD / JAN", "1220005 / 4971635544602", "商品コードとJANコード（バーコード番号）"],
        ],
        col_widths=[5.0, 4.5, 7.0]
    )

    add_para(doc, "")
    add_para(doc, "【予測数値列（左から順）】", bold=True, kwn=True)
    add_sample_table(doc, [
        ["現在庫", "**48**", "現時点の在庫数量（個）。stocks テーブルの合計"],
        ["平均日販", "**3.75**", "過去30日の加重移動平均（WMA）による日次売上数。直近日ほど重みが大きい"],
        ["季節指数", "**1.0**", "今月の季節補正係数。1.0 = 補正なし（通年平均と同じ水準）。夏に需要が増える商品は1.2などになる"],
        ["販促日数", "**0**", "今後30日以内に登録されている販促日数。0 = 販促予定なし"],
        ["販促上振れ", "**0.0**", "販促による追加需要数量（個）。販促なしのため 0"],
        ["受注数量", "**0**", "受注予定（demand_plans）で確定している数量。得意先からの事前注文分"],
        ["30日予測", "**18.0**", "今後30日の需要予測合計（個）。平均日販に曜日指数・季節指数を掛けて30日分積算"],
        ["推奨発注点", "**4**（黄）", "システムが算出した推奨発注点。現行値と異なる場合は黄バッジ表示"],
    ])

    add_note(doc, """なぜ平均日販3.75なのに30日予測が18.0なのか？
3.75は「売上のあった日の平均」です。30日予測は全30日に曜日指数（受注なし曜日は0に近い）を掛けて積算するため、実際に売上が立つ日数分だけの合計になります。
  カバー日数の計算で確認: 48個 ÷（18個÷30日）= 48 ÷ 0.6 = 80.0日 ✓""")

    add_para(doc, "")
    add_para(doc, "【推奨発注点（現行との比較）】", bold=True, kwn=True)
    add_table(doc,
        ["表示", "サンプル", "意味"],
        [
            ["推奨発注点（黄バッジ）", "4", "現行と異なるため黄色で表示。緑バッジは現行と一致している状態"],
            ["現行", "30", "現在 products テーブルに設定されている発注点"],
            ["モードバッジ", "SF", "SF = 安全係数モード / P80 = 80パーセンタイルモード / P90 = 90パーセンタイルモード"],
        ],
        col_widths=[4.5, 2.5, 9.5]
    )

    add_para(doc, "")
    add_para(doc, "【推奨発注数（現行との比較）】", bold=True, kwn=True)
    add_table(doc,
        ["表示", "サンプル", "意味"],
        [
            ["推奨発注数（黄バッジ）", "20", "システムが算出した推奨発注数。リードタイム＋14日分の需要量"],
            ["現行", "40", "現在 products テーブルに設定されている order_qty（基準発注数）"],
        ],
        col_widths=[4.5, 2.5, 9.5]
    )

    add_para(doc, "")
    add_para(doc, "【分析補助データ（右側の列）】", bold=True, kwn=True)
    add_sample_table(doc, [
        ["P80", "**5**", "過去30日の日次売上の80パーセンタイル値。「1日5個以下の日が全体の80%」を意味する"],
        ["P90", "**9**", "90パーセンタイル値。P90モードを選ぶと発注点 = P90 × リードタイム で計算される"],
        ["σ（標準偏差）", "**2.62**", "日次売上のばらつき度合い。値が大きいほど需要変動が激しい"],
        ["補正（✎）", "**✎**（グレー）", "手動調整係数（manual_adj_factor）。✎クリックで商品編集画面へ。現在は1.0（補正なし）"],
        ["カバー日数", "**80.0日**（緑）", "現在庫が何日分もつかの計算値。14日未満で赤・45日超で黄バッジ"],
    ])

    add_note(doc, """この商品（天日干したくあん）のまとめ
・需要は月18個程度（約0.6個/日）と少量
・現行発注点30に対して推奨は4 → 現行が過剰な可能性
・現在庫48個で80日分もある → 過剰在庫の状態
・システムの推奨通りに設定を変更することで在庫適正化が期待できる""")

    # 6.2
    add_heading(doc, "6.2　計算ロジック詳解", 2)

    add_heading(doc, "6.2.1　平均日販（avg_daily）の計算方法", 3)
    add_para(doc, "単純平均ではなく「加重移動平均（WMA）」を使用しています。直近の売上ほど重みが大きく、最新のトレンドを反映します。", kwn=True)
    add_formula(doc, "WMA（加重移動平均）",
        "Σ(売上数量 × 日付順の重み) ÷ Σ(重み)",
        "1日目:2個×1, 2日目:4個×2, 3日目:6個×3 → (2+8+18)÷(1+2+3) = 28÷6 = 4.67個/日")
    add_note(doc, "データが30日分ない場合は過去84日の単純平均にフォールバックします。")

    add_heading(doc, "6.2.2　曜日指数（dow_idx）", 3)
    add_para(doc, "曜日ごとの売上パターンを指数化します。例えば「月曜日だけ受注が集中する」商品では月曜の指数が高くなります。", kwn=True)
    add_table(doc,
        ["曜日", "売上パターン例", "指数"],
        [
            ["月", "受注日（3.75個）", "5.25"],
            ["火", "受注日（3.75個）", "5.25"],
            ["水〜日", "受注なし（0個）", "0.00"],
        ],
        col_widths=[2.5, 5.5, 2.5]
    )
    add_formula(doc, "曜日指数",
        "その曜日の平均売上 ÷ 全日平均日販",
        "月曜平均3.75 ÷ 全日平均0.714 ≈ 5.25")
    add_note(doc, "月・火だけ受注がある場合、30日のうち約9日が受注日 → 予測 ≈ 0.714 × 5.25 × 9 ≈ 33 … 実際の計算は各日ごとに積算します。")

    add_heading(doc, "6.2.3　30日予測の積算方法", 3)
    add_para(doc, "今日から30日間、各日の曜日指数・季節指数・手動調整係数を掛け合わせて積算します。", kwn=True)
    add_formula(doc, "1日分の予測",
        "avg_daily × season_idx × dow_idx × manual_adj_factor",
        "0.714 × 1.0 × 5.25 × 1.0 = 3.75（受注曜日）/ 0.714 × 1.0 × 0.0 × 1.0 = 0（非受注日）")
    add_formula(doc, "30日予測合計",
        "Σ（各日の1日分予測）",
        "3.75 × 5日（月・火が5週） ≈ 18.0個")

    add_heading(doc, "6.2.4　発注点の計算（3モード）", 3)
    add_para(doc, "設定画面で選択したモードにより計算式が変わります。", kwn=True)
    add_table(doc,
        ["モード", "バッジ表示", "計算式", "サンプル計算"],
        [
            ["安全係数モード", "SF",
             "日次予測 × リードタイム × 安全係数",
             "0.6 × 5日 × 1.2 = 3.6 → 切り上げ → 4"],
            ["P80モード", "P80",
             "P80日次値 × リードタイム",
             "5 × 5日 = 25"],
            ["P90モード", "P90",
             "P90日次値 × リードタイム",
             "9 × 5日 = 45"],
        ],
        col_widths=[3.5, 2.5, 5.0, 5.5]
    )
    add_note(doc, "SFモード（安全係数）が最も保守的。P90は欠品を99%防ぐ高いサービスレベルを保証しますが、在庫量は多くなります。")

    add_heading(doc, "6.2.5　推奨発注数の計算", 3)
    add_para(doc, "発注から次回発注までの期間（リードタイム＋2週間）をカバーする数量を推奨します。", kwn=True)
    add_formula(doc, "推奨発注数",
        "日次予測 × (リードタイム + 14日)",
        "0.6 × (5 + 14) = 0.6 × 19 = 11.4 → 切り上げ → 12個")
    add_note(doc, "入数（unit_qty）が設定されている場合はケース単位（入数の倍数）に切り上げます。")

    add_heading(doc, "6.2.6　カバー日数の計算", 3)
    add_para(doc, "現在庫が何日分もつかを示します。在庫過多・過少の判断指標です。", kwn=True)
    add_formula(doc, "カバー日数",
        "現在庫 ÷ 日次予測（= 30日予測 ÷ 30）",
        "48 ÷ (18.0 ÷ 30) = 48 ÷ 0.6 = 80.0日")
    add_table(doc,
        ["表示色", "条件", "意味"],
        [
            ["赤バッジ", "14日未満", "在庫不足の危険域。至急発注が必要"],
            ["緑バッジ", "14〜45日", "適正在庫レンジ"],
            ["黄バッジ", "45日超", "在庫過多。発注点・発注数の見直しを検討"],
        ],
        col_widths=[3.0, 3.5, 10.0]
    )

    # 6.3
    add_heading(doc, "6.3　発注点・発注数の自動反映", 2)
    add_para(doc, "予測画面の「予測値を設定へ反映」機能で、算出した推奨値を商品マスタに一括反映できます。", kwn=True)
    add_step_table(doc, [
        ("反映モードを選択", "「発注点のみ」「発注数のみ」「発注点＋発注数」の3種類から選択"),
        ("検索で絞り込み（任意）", "特定の仕入先や商品だけに絞り込んでから反映することも可能"),
        ("「一括反映」をクリック", "確認ダイアログが表示される"),
        ("確認", "「OK」をクリックすると products テーブルの値が更新される"),
    ])
    add_caution(doc, "一括反映は元に戻せません。事前にCSVエクスポートでバックアップを取ることを推奨します。")

    # 6.4
    add_heading(doc, "6.4　販促予定・受注予定の登録", 2)
    add_para(doc, "販促イベントや得意先からの受注予定を登録すると、30日予測に自動加算されます。")

    add_para(doc, "【販促予定の登録】", bold=True, kwn=True)
    add_sample_table(doc, [
        ["JAN / 商品CD", "4971635544602", "対象商品のJANまたは商品コード（どちらでも可）"],
        ["販促日", "2026-04-10", "販促実施日"],
        ["販促名", "春の特売", "メモ用の名称（任意）"],
        ["上振れ係数", "1.3", "通常需要の1.3倍（30%増し）の需要が見込まれることを意味する"],
    ])
    add_formula(doc, "販促上振れ追加量",
        "日次予測 × (上振れ係数 − 1)",
        "0.6 × (1.3 − 1.0) = 0.18個/日 → 販促日当日に加算")

    add_para(doc, "【受注予定の登録】", bold=True, kwn=True)
    add_sample_table(doc, [
        ["JAN / 商品CD", "4971635544602", "対象商品"],
        ["数量", "100", "受注数量（個）。得意先からの事前注文分"],
        ["受注日", "2026-04-15", "納品日・受注確定日"],
        ["得意先", "○○スーパー", "受注元の得意先名（任意）"],
    ])
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. チェーン・店舗管理
# ══════════════════════════════════════════════════════════════════════════════
def ch07_chains(doc):
    add_heading(doc, "7.　チェーン・店舗マスタ管理", 1)
    add_para(doc, "複数の小売チェーン・店舗を管理します。チェーンごとに商品コード・仕入先コードのマッピングを設定でき、CSV自動インポート時のフィルタリングに使用されます。")

    add_heading(doc, "7.1　マスタ構成", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["マスタ名", "テーブル", "サンプル", "用途"],
        [
            ["チェーンマスタ", "chain_masters", "A社スーパー / チェーンCD: AS01", "チェーンの基本情報と除外フラグ"],
            ["店舗マスタ", "store_masters", "A社錦糸町店 / 店舗CD: AS01-001", "店舗とチェーンの紐付け"],
            ["仕入先コード設定", "supplier_code_settings", "A社チェーン → 仕入先CD: K122", "チェーン別の仕入先コードマッピング"],
            ["商品コード設定", "product_code_settings", "A社チェーン → 商品CD: TK-001", "チェーン別の商品コードマッピング"],
        ],
        col_widths=[3.5, 4.5, 4.5, 4.0]
    )

    add_heading(doc, "7.2　一括インポート", 2)
    add_para(doc, "各マスタとも「Excelテンプレート」をダウンロードして編集後、アップロードして一括登録できます。", kwn=True)
    add_step_table(doc, [
        ("テンプレートDL", "「テンプレートダウンロード」ボタンをクリック"),
        ("編集", "ExcelにチェーンCD・チェーン名などを入力"),
        ("アップロード", "「ファイルを選択」→「インポート」をクリック"),
        ("確認", "登録件数がフラッシュメッセージで表示される"),
    ])
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. CSV自動インポート
# ══════════════════════════════════════════════════════════════════════════════
def ch08_csv(doc):
    add_heading(doc, "8.　CSV自動インポート", 1)
    add_para(doc, "売上データや入荷データをCSVファイルから定期的に自動取り込みします。ネットワーク共有フォルダ（UNCパス）にも対応しています。")

    add_heading(doc, "8.1　インポート設定項目", 2)
    add_para(doc, "設定画面でインポートの条件を登録します。以下のサンプルは「A社の売上CSVを毎日取り込む」設定例です。", kwn=True)
    add_sample_table(doc, [
        ["フォルダパス", r"\\fileserver\sales" + "\\", "CSVの格納フォルダ。UNCパス（ネットワーク共有）も可"],
        ["ファイル名パターン", "{yyyymmdd}_売上実績.csv", "日付プレースホルダを使って当日のファイルを自動特定"],
        ["インポート種別", "sales", "sales=売上（在庫減算）/ receipt=入荷（在庫加算）/ record_only=記録のみ"],
        ["JAN列番号", "2", "CSVの3列目（0始まりなので2）にJANコードがある"],
        ["数量列番号", "5", "CSVの6列目に数量がある"],
        ["エンコード", "Shift-JIS", "UTF-8-SIG / Shift-JIS / CP932 に対応"],
        ["月次モード", "OFF", "OFF=毎日取り込み / ON=月末ファイルを月次として取り込み"],
    ])

    add_heading(doc, "8.2　ファイル名パターンのプレースホルダ", 2)
    add_para(doc, "ファイル名に日付を含むパターンを設定すると、実行日の日付に自動変換されます。", kwn=True)
    add_table(doc,
        ["プレースホルダ", "変換例（2026年4月3日の場合）", "用途"],
        [
            ["{yyyymmdd}", "20260403", "日次ファイル（推奨）"],
            ["{yyyymm}", "202604", "月次ファイル"],
            ["{yymm}", "2604", "2桁年+月"],
            ["{yyyy} / {mm} / {dd}", "2026 / 04 / 03", "分割指定"],
        ],
        col_widths=[4.0, 5.5, 7.0]
    )

    add_heading(doc, "8.3　取り込み処理フロー", 2)
    add_step_table(doc, [
        ("ファイル名解決", r"パターン「{yyyymmdd}_売上実績.csv」 → 「20260403_売上実績.csv」に変換"),
        ("UNC接続", r"\\fileserver へネットワーク接続（認証が必要な場合はnet useで接続）"),
        ("ファイル読み込み", "指定エンコードでCSVを開き、設定したスキップ行数分だけ先頭行を読み飛ばす"),
        ("行フィルタリング", "部門コード・チェーンCDのフィルタ設定がある場合は条件に合う行のみ処理"),
        ("重複チェック", "各行をSHA256ハッシュで識別。既に取り込み済みの行はスキップ（2重取込防止）"),
        ("JAN正規化", "指数表記（4.9E+12など）を自動変換。例: 4.9E+12 → 4900000000000"),
        ("在庫更新", "sales: stocks.quantity を減算 / receipt: _record_receipt() で在庫追加"),
        ("集計テーブル更新", "sales_daily_agg に日次集計データをUPSERT。予測キャッシュを自動無効化"),
    ])
    add_caution(doc, "ファイルが見つからない場合、または形式エラーの行がある場合はログファイルにエラーが記録されます。logs/stderr.log を確認してください。")
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. メール通知
# ══════════════════════════════════════════════════════════════════════════════
def ch09_mail(doc):
    add_heading(doc, "9.　メール通知", 1)

    add_heading(doc, "9.1　発注書メール", 2)
    add_para(doc, "手動発注または自動チェック発注時に、仕入先ごとに1通の発注書メールが自動送信されます。", kwn=True)
    add_sample_table(doc, [
        ["送信先", "order@kimura-tsukemono.co.jp", "商品マスタの supplier_email に設定したアドレス"],
        ["件名", "【発注書】キムラ漬物 2026-04-03", "「【発注書】仕入先名 発注日」の形式"],
        ["本文", "商品CD: 1220005 / 天日干したくあん / 42個", "商品コード・商品名・発注数量の一覧"],
        ["送信方式", "仕入先1社に1通まとめて送信", "複数商品を1社に発注する場合でも1通にまとめる"],
        ["送信結果", "送信済み", "order_history.mail_sent に記録。失敗時は mail_result にエラー内容"],
    ])
    add_note(doc, "送信失敗した場合は order_history の「再送信」ボタンから再送できます。")

    add_heading(doc, "9.2　賞味期限アラートメール", 2)
    add_para(doc, "スケジューラー実行時に、賞味期限が近い在庫をシステム管理者にメール通知します。", kwn=True)
    add_sample_table(doc, [
        ["検出条件", "残り30日以内", "商品マスタの expiry_alert_days（デフォルト30日）以内"],
        ["送信頻度", "1日1回", "スケジューラー実行時に1通送信"],
        ["通知内容", "天日干したくあん / 在庫12個 / 2026-04-15 期限切れ", "アラート対象商品のJAN・商品名・在庫数・賞味期限・ロケーション"],
    ])
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 10. 権限管理
# ══════════════════════════════════════════════════════════════════════════════
def ch10_permissions(doc):
    add_heading(doc, "10.　権限管理", 1)

    add_heading(doc, "10.1　権限の種類と対象画面", 2)
    add_para(doc, "ユーザーごとにアクセスできる画面を個別に設定します。管理者のみが権限設定を変更できます。", kwn=True)
    add_table(doc,
        ["権限キー", "アクセス対象", "設定例"],
        [
            ["dashboard", "ダッシュボード（トップページ）", "全ユーザーに付与"],
            ["inventory", "在庫一覧・編集・廃棄・移動", "倉庫担当者"],
            ["products", "商品マスタ登録・編集・インポート", "商品担当者"],
            ["orders", "発注チェック・手動発注・発注残管理", "発注担当者"],
            ["reports", "需要予測・ABC分析・52週MD", "分析担当者"],
            ["chains", "チェーン・店舗マスタ管理", "マスタ管理担当"],
            ["admin", "ユーザー管理・システム設定（全権限相当）", "システム管理者のみ"],
        ],
        col_widths=[3.5, 6.0, 7.0]
    )

    add_heading(doc, "10.2　権限設定の方法", 2)
    add_step_table(doc, [
        ("管理者でログイン", "admin 権限を持つユーザーでログイン"),
        ("ユーザー管理を開く", "メニューの「設定」→「ユーザー管理」をクリック"),
        ("対象ユーザーを選択", "権限を変更したいユーザーの「編集」ボタンをクリック"),
        ("権限チェック", "付与する権限のチェックボックスをオン/オフ"),
        ("保存", "「保存」ボタンをクリック。次回ログイン時から反映"),
    ])
    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 付録. 主要テーブル一覧
# ══════════════════════════════════════════════════════════════════════════════
def appendix_tables(doc):
    add_heading(doc, "付録.　主要テーブル一覧", 1)
    add_para(doc, "システムが使用するデータベースの主要テーブルと用途の一覧です。", kwn=True)
    add_table(doc,
        ["テーブル名", "用途", "主なカラム"],
        [
            ["products", "商品マスタ", "jan, product_cd, product_name, supplier_cd, reorder_point, reorder_auto, lead_time_days"],
            ["stocks", "在庫台帳（ロット単位）", "jan, quantity, expiry_date, lot_no, location_code"],
            ["sales_history", "売上履歴（CSVインポートデータ）", "jan, quantity, sale_date, source_file, row_hash"],
            ["sales_daily_agg", "日次売上集計（予測キャッシュ）", "jan, sale_date, total_qty"],
            ["order_history", "発注履歴", "jan, order_qty, order_date, expected_receipt_date, mail_sent, trigger_type"],
            ["order_pending", "混載保留中の発注", "jan, order_qty, mixed_group, status, force_send_date"],
            ["disposed_stocks", "廃棄済み在庫", "jan, quantity, reason_type, loss_amount, disposed_at"],
            ["stock_movements", "在庫変動ログ（全操作の監査証跡）", "jan, move_type, quantity, before_qty, after_qty, note"],
            ["stock_transfers", "ロケーション間移動ログ", "from_stock_id, to_stock_id, from_location, to_location"],
            ["csv_import_settings", "CSVインポート設定", "folder_path, filename_pattern, col_jan, col_qty, import_type, encoding"],
            ["promotion_plans", "販促計画（予測の上振れ反映）", "jan, promo_date, promo_name, uplift_factor"],
            ["demand_plans", "受注予定（確定需要）", "jan, demand_date, demand_qty, customer_name"],
            ["chain_masters", "チェーンマスタ", "chain_cd, chain_name, is_excluded"],
            ["store_masters", "店舗マスタ", "store_cd, store_name, chain_cd"],
            ["alert_logs", "アラートイベントログ", "jan, alert_type, triggered_at, message"],
            ["users", "ユーザーマスタ", "username, password_hash, permissions, must_change_password"],
            ["settings", "システム設定", "key, value（forecast_ai_mode, safety_level_z など）"],
            ["temp_sensitivity", "気温感応度（問屋向け予測）", "jan, temp_coef, r_squared, base_temp"],
            ["weekly_md_plans", "52週MDプラン", "jan, week_start, plan_qty, actual_qty"],
            ["weather_data", "気温データ", "obs_date, avg_temp, location"],
        ],
        col_widths=[4.5, 4.0, 8.0]
    )

# ══════════════════════════════════════════════════════════════════════════════
# メイン
# ══════════════════════════════════════════════════════════════════════════════
def main():
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = '游ゴシック'
    doc.styles['Normal'].font.size = Pt(10)

    build_cover(doc)
    build_toc(doc)
    ch01_login(doc)
    ch02_dashboard(doc)
    ch03_inventory(doc)
    ch04_products(doc)
    ch05_orders(doc)
    ch06_forecast(doc)
    ch07_chains(doc)
    ch08_csv(doc)
    ch09_mail(doc)
    ch10_permissions(doc)
    appendix_tables(doc)

    doc.save(OUT)
    print(f"生成完了: {OUT}")

if __name__ == '__main__':
    main()
