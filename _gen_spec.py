# -*- coding: utf-8 -*-
"""在庫管理システム 要件定義・仕様設計書 生成スクリプト"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "在庫管理システム_仕様設計書.docx"

# ─── XML ヘルパー ────────────────────────────────────────────────────────────
def _kwn(p):
    pPr = p._p.get_or_add_pPr()
    kn = OxmlElement('w:keepNext')
    pPr.append(kn)

def _kt(p):
    pPr = p._p.get_or_add_pPr()
    kt = OxmlElement('w:keepLines')
    pPr.append(kt)

def _cant_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit')
    cs.set(qn('w:val'), '1')
    trPr.append(cs)

def _row_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), '1')
    trPr.append(th)

def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_col_widths(table, widths_cm):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])

# ─── ドキュメント ヘルパー ────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    _kwn(p)
    _kt(p)
    return p

def add_para(doc, text, kwn=False, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    _kt(p)
    if kwn:
        _kwn(p)
    return p

def add_note(doc, text):
    """注意書き（グレー小字）"""
    p = doc.add_paragraph()
    r = p.add_run(f"※ {text}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _kt(p)
    return p

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    p_anchor = doc.add_paragraph()
    p_anchor.paragraph_format.space_before = Pt(0)
    p_anchor.paragraph_format.space_after = Pt(0)
    # run だけ削除し pPr は残す
    from lxml import etree
    for child in list(p_anchor._p):
        if child.tag != qn('w:pPr'):
            p_anchor._p.remove(child)
    _kwn(p_anchor)  # pPr クリア後に追加

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダー行
    hdr = tbl.rows[0]
    _row_header(hdr)
    _cant_split(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        _shade_cell(c, header_color)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # データ行
    for i, row_data in enumerate(rows):
        row = tbl.rows[i + 1]
        _cant_split(row)
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            c = row.cells[j]
            _shade_cell(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)

    if col_widths:
        _set_col_widths(tbl, col_widths)
    return tbl

def add_step_table(doc, steps):
    """番号付きステップ表（ステップ番号 | タイトル | 説明）"""
    headers = ["STEP", "処理名", "処理内容"]
    col_widths = [1.2, 3.5, 11.8]
    tbl = add_table(doc, headers, [], col_widths=col_widths, header_color="2E75B6")
    for i, (title, desc) in enumerate(steps, 1):
        row = tbl.add_row()
        _cant_split(row)
        bg = "EBF3FB" if i % 2 == 1 else "FFFFFF"
        data = [f"⑩"[:1] + str(i) if i > 9 else ["①","②","③","④","⑤","⑥","⑦","⑧","⑨","⑩"][i-1], title, desc]
        for j, val in enumerate(data):
            c = row.cells[j]
            _shade_cell(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if j == 0:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.bold = True
    return tbl

def page_break(doc):
    doc.add_page_break()

# ─── 表紙 ────────────────────────────────────────────────────────────────────
def build_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph("在庫管理システム")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(28)
    t.runs[0].bold = True
    t.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    t2 = doc.add_paragraph("要件定義・機能仕様設計書")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(20)
    t2.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    for _ in range(3):
        doc.add_paragraph()
    t3 = doc.add_paragraph("Version 1.0　　2026年4月")
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.runs[0].font.size = Pt(12)
    t3.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    page_break(doc)

# ─── 目次（静的） ─────────────────────────────────────────────────────────────
def build_toc(doc):
    add_heading(doc, "目次", 1)
    toc = [
        ("1", "システム概要"),
        ("2", "認証・ログイン管理"),
        ("3", "ダッシュボード"),
        ("4", "在庫管理"),
        ("4.1", "在庫一覧・検索"),
        ("4.2", "在庫直接編集"),
        ("4.3", "在庫廃棄・復元"),
        ("4.4", "ロケーション移動"),
        ("4.5", "ピッキングプラン生成"),
        ("4.6", "補充タスク管理"),
        ("5", "商品マスタ管理"),
        ("5.1", "商品CRUD"),
        ("5.2", "一括インポート / エクスポート"),
        ("6", "発注管理"),
        ("6.1", "手動発注"),
        ("6.2", "自動チェック実行"),
        ("6.3", "混載グループ管理"),
        ("6.4", "発注残・部分入荷管理"),
        ("7", "需要予測"),
        ("7.1", "AIモード予測"),
        ("7.2", "前年実績モード予測"),
        ("8", "チェーン・店舗マスタ管理"),
        ("9", "CSVインポート自動処理"),
        ("10", "メール通知"),
        ("11", "データベース設計（主要テーブル）"),
        ("12", "権限管理"),
    ]
    for num, title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        indent = 0 if '.' not in num else 1
        p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(f"{num}　{title}")
        r.font.size = Pt(10)
        if '.' not in num:
            r.bold = True
    page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# 各章
# ═══════════════════════════════════════════════════════════════════════════════

def ch01_overview(doc):
    add_heading(doc, "1.　システム概要", 1)
    add_para(doc, "本システムは、小売チェーン向けの在庫管理・自動発注・需要予測を統合したWebアプリケーションです。Flask (Python) をバックエンドに、PostgreSQLをデータベースとして使用します。")
    add_para(doc, "")

    add_heading(doc, "1.1　主要機能一覧", 2)
    add_table(doc,
        ["機能カテゴリ", "主な機能"],
        [
            ["認証・権限管理", "ログイン / ログアウト、初回パスワード変更強制、ロールベースのページ権限"],
            ["ダッシュボード", "在庫状況サマリ、発注アラート、賞味期限アラート、直近発注履歴"],
            ["在庫管理", "在庫一覧・検索・編集、廃棄・復元、ロケーション移動、ピッキング・補充"],
            ["商品マスタ", "商品CRUD、Excelテンプレート、一括インポート/エクスポート"],
            ["発注管理", "手動発注、自動チェック、混載グループ発注、部分入荷・発注残管理"],
            ["需要予測", "AIモード（機械学習的推計）/ 前年実績モード（前年同期比）"],
            ["チェーン・店舗管理", "チェーンマスタ、店舗マスタ、仕入先コード・商品コード紐付け"],
            ["CSVインポート", "売上・入荷CSVの定期自動取り込み、UNCパス対応"],
            ["メール通知", "発注書メール、賞味期限アラートメール、仕入先ごと一括送信"],
        ],
        col_widths=[4.5, 12.0]
    )

    add_heading(doc, "1.2　技術スタック", 2)
    add_table(doc,
        ["区分", "採用技術"],
        [
            ["Webフレームワーク", "Flask (Python 3.x)"],
            ["データベース", "PostgreSQL"],
            ["認証", "Flask-Login + セッション管理 + レート制限（Flask-Limiter）"],
            ["メール送信", "smtplib / mail_service モジュール"],
            ["Excel出力", "openpyxl"],
            ["Word文書生成", "python-docx"],
            ["スケジューラー", "threading.Timer（バックグラウンドスレッド）"],
            ["Windowsサービス", "NSSM（Non-Sucking Service Manager）"],
            ["CSRFガード", "Flask-WTF"],
        ],
        col_widths=[4.5, 12.0]
    )
    page_break(doc)


def ch02_auth(doc):
    add_heading(doc, "2.　認証・ログイン管理", 1)

    add_heading(doc, "2.1　機能概要", 2)
    add_para(doc, "ユーザー認証とページごとの権限管理を行います。ログイン失敗のレート制限、初回パスワード変更の強制、ロールベースのアクセス制御（RBAC）を実装しています。", kwn=True)

    add_heading(doc, "2.2　ログイン処理フロー", 2)
    add_step_table(doc, [
        ("ログインフォーム表示", "GET /login でログイン画面を表示。既にセッションが有効な場合はダッシュボードへリダイレクト。"),
        ("認証情報検証", "POST /login でユーザー名・パスワードを受け取り、DBの users テーブルと照合（bcrypt ハッシュ比較）。"),
        ("レート制限チェック", "同一IPからのログイン試行を10回/分に制限（Flask-Limiter）。超過時は429エラーを返す。"),
        ("初回パスワード確認", "must_change_password = 1 の場合はパスワード変更画面にリダイレクト。通常ログインは不可。"),
        ("セッション確立", "認証成功時にFlaskセッションにユーザーID・権限情報を格納し、ダッシュボードへリダイレクト。"),
        ("ログアウト", "GET /logout でセッションをクリアしてログイン画面へ戻す。"),
    ])

    add_heading(doc, "2.3　権限（パーミッション）一覧", 2)
    add_table(doc,
        ["権限キー", "アクセス対象画面"],
        [
            ["dashboard", "ダッシュボード"],
            ["inventory", "在庫管理（一覧・編集・廃棄・移動）"],
            ["products", "商品マスタ管理"],
            ["orders", "発注管理（手動発注・発注チェック）"],
            ["forecast", "需要予測"],
            ["chains", "チェーン・店舗マスタ管理"],
            ["admin", "ユーザー管理・システム設定（全機能に相当）"],
        ],
        col_widths=[4.5, 12.0]
    )

    add_heading(doc, "2.4　制約・注意点", 2)
    add_note(doc, "パスワードは bcrypt でハッシュ化して保存します。平文保存は禁止です。")
    add_note(doc, "レート制限はメモリベースのため、サーバー再起動でリセットされます。")
    add_note(doc, "must_change_password フラグはシステム管理者が手動でOFF（0に更新）することも可能です。")
    page_break(doc)


def ch03_dashboard(doc):
    add_heading(doc, "3.　ダッシュボード", 1)

    add_heading(doc, "3.1　機能概要", 2)
    add_para(doc, "ログイン直後に表示されるトップページ。在庫・発注・賞味期限の状態をリアルタイムに把握できるサマリ画面です。", kwn=True)

    add_heading(doc, "3.2　表示項目", 2)
    add_table(doc,
        ["表示区分", "内容", "データソース"],
        [
            ["有効商品数", "is_active=1 の商品総数", "products"],
            ["発注アラート数", "発注点以下かつ未発注の商品数", "products + stocks"],
            ["賞味期限アラート数", "30日以内に期限切れとなる在庫のロット数", "stocks"],
            ["直近発注履歴", "直近10件の発注（日付・商品名・数量・メール結果）", "order_history"],
            ["アラートログ", "直近のアラートイベント（発注点到達・ロット到達など）", "alert_logs"],
        ],
        col_widths=[4.0, 7.0, 5.5]
    )

    add_heading(doc, "3.3　発注アラートの判定ロジック", 2)
    add_para(doc, "以下の条件をすべて満たす商品が「発注アラート」としてカウントされます。", kwn=True)
    add_table(doc,
        ["条件", "内容"],
        [
            ["is_active = 1", "有効商品であること"],
            ["reorder_point > 0", "発注点が設定されていること"],
            ["現在庫 ≤ 発注点", "在庫数量が発注点以下であること"],
            ["ordered_at IS NULL", "当日未発注であること（本日の発注フラグなし）"],
        ],
        col_widths=[5.0, 11.5]
    )
    page_break(doc)


def ch04_inventory(doc):
    add_heading(doc, "4.　在庫管理", 1)

    # 4.1
    add_heading(doc, "4.1　在庫一覧・検索", 2)
    add_para(doc, "現在の在庫をロット・ロケーション単位で一覧表示します。JAN・商品コード・商品名・仕入先による絞り込み検索が可能です。", kwn=True)
    add_table(doc,
        ["項目", "内容"],
        [
            ["表示単位", "stocks テーブルの1レコード = 1ロット"],
            ["表示列", "JAN / 商品名 / 仕入先 / ロケーション / 数量 / 賞味期限 / ロット番号"],
            ["検索対象", "JAN・商品CD・商品名・仕入先CD・仕入先名（部分一致）"],
            ["エクスポート", "CSV / Excel（.xlsx）形式で全件ダウンロード可能"],
            ["ソート", "仕入先CD数値順 → 商品CD数値順"],
        ],
        col_widths=[4.0, 12.5]
    )

    # 4.2
    add_heading(doc, "4.2　在庫直接編集", 2)
    add_para(doc, "在庫数量・賞味期限・ロケーションを直接変更します。変更履歴はstock_movementsに記録されます。", kwn=True)
    add_step_table(doc, [
        ("編集フォーム表示", "在庫一覧から対象レコードを選択し編集画面を開く。"),
        ("変更前数量の記録", "現在の quantity を before_qty として保持。"),
        ("在庫レコード更新", "stocks テーブルの quantity / expiry_date / location_code を更新。"),
        ("移動履歴記録", "stock_movements に move_type='edit'、変更前後数量を INSERT。"),
        ("コミット", "DB コミット後、在庫一覧画面へリダイレクト。"),
    ])

    # 4.3
    add_heading(doc, "4.3　在庫廃棄・復元", 2)
    add_para(doc, "期限切れ・破損などの理由で在庫を廃棄します。廃棄レコードはdisposed_stocksに保存され、誤廃棄の場合は復元できます。", kwn=True)
    add_table(doc,
        ["区分", "内容"],
        [
            ["廃棄理由コード", "expiry（賞味期限）/ damage（破損）/ other（その他）"],
            ["廃棄ロス金額", "廃棄数量 × 原価（cost_price）を自動計算してDBに保存"],
            ["廃棄後処理", "stocks レコードを削除し disposed_stocks に移動。stock_movements に廃棄ログ記録"],
            ["復元", "disposed_stocks から stocks へ再インサート。disposed フラグを更新"],
        ],
        col_widths=[4.0, 12.5]
    )

    # 4.4
    add_heading(doc, "4.4　ロケーション移動", 2)
    add_para(doc, "在庫を別のロケーション（保管場所）へ移動します。移動元・移動先の両レコードと移動ログが記録されます。", kwn=True)
    add_step_table(doc, [
        ("移動フォーム入力", "移動元在庫ID・移動先ロケーション・移動数量を入力。"),
        ("数量チェック", "移動数量 > 移動元在庫数量の場合はエラー返却。"),
        ("移動元在庫更新", "移動元の quantity を減算。0になった場合はレコード削除。"),
        ("移動先在庫更新", "移動先ロケーションに同一JAN・同一賞味期限のレコードがあれば加算、なければ新規INSERT。"),
        ("移動ログ記録", "stock_transfers に from_stock_id / to_stock_id / from_location / to_location を記録。"),
    ])

    # 4.5
    add_heading(doc, "4.5　ピッキングプラン生成", 2)
    add_para(doc, "今後7日間の出荷計画に基づき、ロケーション・賞味期限を考慮した最適なピッキング順序を生成します。", kwn=True)
    add_table(doc,
        ["パラメータ", "デフォルト値", "説明"],
        [
            ["予測ホライズン", "7日", "何日分の需要を対象にするか"],
            ["ピッキング優先順", "賞味期限 昇順", "古い在庫から先にピッキング（FEFO）"],
            ["出力形式", "画面表示 / Excel", "商品・ロケーション・数量・賞味期限のリスト"],
        ],
        col_widths=[4.0, 3.0, 9.5]
    )

    # 4.6
    add_heading(doc, "4.6　補充タスク管理", 2)
    add_para(doc, "バックヤード在庫を売り場へ補充するタスクを管理します。補充が必要な商品を自動算出し、担当者への作業指示書として機能します。", kwn=True)
    add_table(doc,
        ["機能", "説明"],
        [
            ["補充対象算出", "売り場在庫 < 補充閾値 の商品をリストアップ"],
            ["タスク登録", "補充数量・担当者・期限を replenishment_history に記録"],
            ["完了処理", "担当者が完了登録すると在庫が移動（stocks 更新）"],
            ["Excelエクスポート", "補充作業リストをExcel出力可能"],
        ],
        col_widths=[4.0, 12.5]
    )
    page_break(doc)


def ch05_products(doc):
    add_heading(doc, "5.　商品マスタ管理", 1)

    add_heading(doc, "5.1　機能概要", 2)
    add_para(doc, "商品の登録・編集・削除・一括インポート/エクスポートを管理します。発注に必要なすべてのパラメータをここで設定します。")

    add_heading(doc, "5.2　商品マスタ項目一覧", 2)
    add_para(doc, "商品マスタ（products テーブル）の主要カラムは以下のとおりです。", kwn=True)
    add_table(doc,
        ["カラム名", "型", "説明"],
        [
            ["jan", "TEXT", "JANコード（主キー）。指数表記・小数点付きを正規化して保存"],
            ["product_cd", "TEXT", "商品コード（仕入先管理番号）"],
            ["product_name", "TEXT", "商品名"],
            ["supplier_cd", "TEXT", "仕入先コード"],
            ["supplier_name", "TEXT", "仕入先名"],
            ["supplier_email", "TEXT", "仕入先メールアドレス（発注書送信先）"],
            ["unit_qty", "INTEGER", "1ケース入数（発注はケース単位に切り上げ）"],
            ["order_qty", "INTEGER", "標準発注数量（ケース数）"],
            ["reorder_point", "INTEGER", "発注点（この在庫数以下で発注トリガー）"],
            ["reorder_auto", "INTEGER", "発注モード: 0=手動 / 1=AI / 2=前年実績"],
            ["lead_time_days", "INTEGER", "リードタイム（発注から入荷までの日数）"],
            ["safety_factor", "REAL", "安全係数（需要予測に掛ける係数）"],
            ["shelf_life_days", "INTEGER", "賞味期限日数（入荷日からの日数）"],
            ["expiry_alert_days", "INTEGER", "賞味期限アラート閾値（残り日数）"],
            ["mixed_group", "TEXT", "混載グループ名（同グループで合計ロット数を管理）"],
            ["mixed_force_days", "INTEGER", "混載強制送信日数（未達でもN日後に強制送信）"],
            ["cost_price", "REAL", "原価（廃棄ロス計算に使用）"],
            ["sell_price", "REAL", "販売価格"],
            ["is_active", "INTEGER", "有効フラグ（1=有効 / 0=無効・論理削除）"],
            ["ordered_at", "DATE", "本日発注済みフラグ（当日発注後にセット、翌日リセット）"],
            ["location_code", "TEXT", "デフォルトロケーションコード"],
        ],
        col_widths=[4.0, 2.0, 10.5]
    )

    add_heading(doc, "5.3　発注モード（reorder_auto）", 2)
    add_para(doc, "reorder_auto カラムの値によって自動発注の動作が変わります。", kwn=True)
    add_table(doc,
        ["値", "モード名", "動作"],
        [
            ["0", "手動", "自動チェックの対象外。手動発注のみ"],
            ["1", "AIモード", "過去売上データから機械学習的に発注点を自動算出・更新"],
            ["2", "前年実績モード", "前年同期の売上実績をベースに発注点を算出"],
        ],
        col_widths=[1.5, 4.0, 11.0]
    )

    add_heading(doc, "5.4　一括インポート仕様", 2)
    add_para(doc, "Excel・CSV形式での商品マスタ一括登録・更新に対応しています。", kwn=True)
    add_table(doc,
        ["項目", "仕様"],
        [
            ["対応形式", "CSV（UTF-8 / Shift-JIS）/ Excel（.xlsx）"],
            ["インポートモード", "upsert：既存商品は更新、新規は追加 / add_only：新規追加のみ"],
            ["JANコード正規化", "指数表記（4.9E+12）や小数点付きを自動変換"],
            ["日付正規化", "YYYY/MM/DD・YYYYMMDD・YY/MM/DD など10種類以上のフォーマットを自動判定"],
            ["エラー処理", "行単位でスキップ。エラー行一覧をフラッシュメッセージで通知"],
            ["Excelテンプレート", "ダウンロード可能。発注モード列にドロップダウン入力規則付き"],
        ],
        col_widths=[4.0, 12.5]
    )
    page_break(doc)


def ch06_orders(doc):
    add_heading(doc, "6.　発注管理", 1)

    # 6.1
    add_heading(doc, "6.1　手動発注", 2)
    add_para(doc, "担当者が画面上で商品を選択し、任意のタイミングで発注を行います。", kwn=True)
    add_step_table(doc, [
        ("商品選択", "発注チェック画面の商品一覧から発注したい商品のチェックボックスを選択。"),
        ("数量確認", "order_qty（標準発注数量）が初期表示される。必要に応じて変更可能。"),
        ("ケース単位丸め", "unit_qty（入数）が2以上の場合、入力数量を1ケース単位に切り上げ（ceiling計算）。"),
        ("order_history 記録", "発注日・商品情報・数量・trigger_type='manual' を INSERT。"),
        ("メールキュー追加", "queue_order() で仕入先ごとのメールキューに追加。"),
        ("一括メール送信", "flush_order_mail() で仕入先ごとに1通まとめて発注書メール送信。"),
        ("発注済みフラグ更新", "products.ordered_at に本日日付をセット（当日重複防止）。"),
        ("order_pending 削除", "当該商品の保留レコードがあれば削除。"),
    ])
    add_note(doc, "手動発注と自動チェック発注は run_order_check() を共有しているため動作は同一です。")

    # 6.2
    add_heading(doc, "6.2　自動チェック実行", 2)
    add_para(doc, "発注点・ロット数の条件を自動判定し、条件を満たした商品を自動発注します。", kwn=True)

    add_table(doc,
        ["トリガー方法", "タイミング"],
        [
            ["手動", "発注チェック画面の「🔄 自動チェック実行」ボタン"],
            ["自動（スケジューラー）", "毎日設定時刻（デフォルト 朝8:00）にバックグラウンドスレッドが実行"],
        ],
        col_widths=[5.0, 11.5]
    )

    add_para(doc, "")
    add_para(doc, "【処理フロー】", bold=True, kwn=True)
    add_step_table(doc, [
        ("対象商品の抽出", "is_active=1 かつ reorder_auto IN (1,2) の全商品をループ処理。"),
        ("有効在庫の計算", "有効在庫 = 現在庫 − expiry_alert_days 以内の期限切れ間近在庫（販売不可分を除外）。"),
        ("当日・保留中の重複チェック", "order_history に当日レコードあり → スキップ。order_pending に保留中レコードあり → スキップ。"),
        ("発注トリガーの判定", "発注点到達：有効在庫 ≤ 発注点 / ロット数到達：在庫 ≥ メーカーロット数。"),
        ("発注数量の計算", "不足数 = 発注点 − 有効在庫 / 発注ケース数 = ⌈不足数 ÷ 発注数量⌉（1ケース単位切り上げ）。"),
        ("混載グループ判定", "混載グループあり → order_pending に保留登録（強制送信日=本日+mixed_force_days）。混載グループなし → 即時発注・メールキューへ追加。"),
        ("混載条件チェック（_check_mixed_groups）", "gte モード：グループ合計 ≥ 混載ケース数 / unit モード：合計が混載ケース数の倍数 / 強制送信日超過は条件未達でも送信。"),
        ("メール一括送信", "仕入先ごとに1通の発注書メールを送信。"),
        ("DB記録更新", "order_history・order_pending・products.ordered_at・alert_logs を更新。"),
    ])

    add_para(doc, "")
    add_para(doc, "【発注チェック画面の3区分】", bold=True, kwn=True)
    add_table(doc,
        ["区分", "表示内容"],
        [
            ["発注対象商品", "発注点以下で今すぐ発注が必要な商品"],
            ["混載ロット保留中", "条件待ちで order_pending に登録済みの商品"],
            ["発注済み", "本日すでに発注完了・納品待ちの商品"],
        ],
        col_widths=[5.0, 11.5]
    )
    add_note(doc, "1日1回制限：同一商品は ordered_at で管理し、当日中は再発注されません。")
    add_note(doc, "混載強制送信：mixed_force_days 日後は混載ケース数未達でも自動送信します。")

    # 6.3
    add_heading(doc, "6.3　混載グループ管理", 2)
    add_para(doc, "複数商品をまとめて1ロット単位で仕入先に発注する「混載発注」の管理機能です。", kwn=True)
    add_table(doc,
        ["設定項目", "説明"],
        [
            ["mixed_group", "グループ名（同じ名前の商品が同一グループ）"],
            ["mixed_lot_qty", "グループの合計混載ケース数閾値"],
            ["mixed_mode", "判定モード：gte（以上）/ unit（倍数）"],
            ["mixed_force_days", "強制送信日数（保留開始から N 日後に強制送信）"],
        ],
        col_widths=[4.0, 12.5]
    )
    add_para(doc, "")
    add_para(doc, "画面操作：", bold=True, kwn=True)
    add_table(doc,
        ["操作", "内容"],
        [
            ["数量調整", "グループ内各商品の発注数量を手動で変更可能"],
            ["個別強制送信", "1商品だけを即時送信（グループ条件を無視）"],
            ["グループ一括強制送信", "グループ全体を条件無視で即時送信"],
        ],
        col_widths=[4.0, 12.5]
    )

    # 6.4
    add_heading(doc, "6.4　発注残・部分入荷管理", 2)
    add_para(doc, "発注済み商品の入荷状況を管理します。部分入荷（分割納品）と未入荷残量（バックオーダー）をトラッキングします。", kwn=True)
    add_table(doc,
        ["機能", "説明"],
        [
            ["発注残一覧", "order_history で mail_sent=1 かつ未入荷（received_qty < order_qty）の商品を表示"],
            ["遅延分類", "過去の expected_receipt_date を超えた発注は「延滞」、近い場合は「要注意」と色分け表示"],
            ["受領登録", "入荷数量・賞味期限・ロット番号・ロケーションを入力して stocks に追加。stock_movements にも記録"],
            ["部分入荷", "order_qty に対して received_qty が少ない場合は「発注残」として引き続き管理"],
            ["クローズ処理", "完全入荷またはキャンセル理由を入力してクローズ。order_history に completed_at をセット"],
        ],
        col_widths=[4.0, 12.5]
    )
    add_note(doc, "受領登録は _record_receipt() 関数が担当し、stocks テーブルへのINSERTと stock_movements のログ記録を一括処理します。")
    page_break(doc)


def ch07_forecast(doc):
    add_heading(doc, "7.　需要予測", 1)

    add_heading(doc, "7.1　機能概要", 2)
    add_para(doc, "過去の売上データを分析し、将来の需要を予測して発注点を自動更新します。AIモードと前年実績モードの2種類があります。")

    add_heading(doc, "7.2　AIモード予測", 2)
    add_para(doc, "過去の sales_history データをベースに、トレンド・季節性・プロモーション計画を加味した予測を行います。", kwn=True)
    add_table(doc,
        ["処理項目", "内容"],
        [
            ["データソース", "sales_history（過去売上）/ promotion_plans（販促計画）/ demand_plans（需要計画）"],
            ["集計単位", "sales_daily_agg テーブルに日次集計済みデータをキャッシュして使用"],
            ["予測ロジック", "_build_forecast_rows() で移動平均・トレンド補正・安全係数（safety_factor）を適用"],
            ["発注点更新", "予測需要 × リードタイム日数 + 安全在庫 = 新発注点として products.reorder_point を更新"],
            ["キャッシュ", "予測結果は24時間TTLでメモリキャッシュ。バックグラウンドで自動リフレッシュ"],
        ],
        col_widths=[4.0, 12.5]
    )

    add_heading(doc, "7.3　前年実績モード予測", 2)
    add_para(doc, "前年の同週・同月の売上実績をそのまま今期の予測値として使用します。季節商品や安定品に適しています。", kwn=True)
    add_table(doc,
        ["処理項目", "内容"],
        [
            ["比較期間", "前年同週（±3日）の売上実績"],
            ["補正係数", "safety_factor を乗算して安全在庫を加算"],
            ["発注点更新", "前年実績 × safety_factor × リードタイム日数 = 新発注点"],
            ["月次自動更新", "auto_check.py のスケジューラーが月初に全商品を一括更新"],
        ],
        col_widths=[4.0, 12.5]
    )

    add_heading(doc, "7.4　予測画面の表示内容", 2)
    add_table(doc,
        ["表示列", "内容"],
        [
            ["商品名 / JAN", "対象商品の基本情報"],
            ["現在庫", "現時点の実在庫数量"],
            ["現発注点", "現在設定されている発注点"],
            ["予測需要（7日）", "今後7日間の予測需要数量"],
            ["予測需要（30日）", "今後30日間の予測需要数量"],
            ["推奨発注点", "システムが算出した推奨発注点"],
            ["乖離", "現発注点と推奨発注点の差"],
        ],
        col_widths=[4.0, 12.5]
    )
    page_break(doc)


def ch08_chains(doc):
    add_heading(doc, "8.　チェーン・店舗マスタ管理", 1)

    add_heading(doc, "8.1　機能概要", 2)
    add_para(doc, "複数の小売チェーン・店舗を管理します。チェーンごとに仕入先コード・商品コードのマッピングを設定でき、CSVインポート時のフィルタリングや除外設定に使用されます。")

    add_heading(doc, "8.2　マスタ構成", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["マスタ名", "テーブル", "主な管理内容"],
        [
            ["チェーンマスタ", "chain_masters", "チェーンCD・チェーン名・除外フラグ"],
            ["店舗マスタ", "store_masters", "店舗CD・店舗名・所属チェーンCD"],
            ["仕入先コード設定", "supplier_code_settings", "チェーン/店舗ごとの仕入先コードマッピング"],
            ["商品コード設定", "product_code_settings", "チェーン/店舗ごとの商品コードマッピング"],
        ],
        col_widths=[4.0, 4.5, 8.0]
    )

    add_heading(doc, "8.3　一括インポート", 2)
    add_table(doc,
        ["対象", "操作"],
        [
            ["チェーンマスタ", "Excelテンプレートをダウンロードして編集後、一括アップロード"],
            ["店舗マスタ", "同上"],
            ["仕入先コード設定", "同上"],
            ["商品コード設定", "同上"],
        ],
        col_widths=[5.0, 11.5]
    )
    page_break(doc)


def ch09_csv_import(doc):
    add_heading(doc, "9.　CSVインポート自動処理", 1)

    add_heading(doc, "9.1　機能概要", 2)
    add_para(doc, "売上データ・入荷データをCSVファイルから自動取り込みします。ネットワーク共有フォルダ（UNCパス）にも対応し、スケジューラーが設定時刻に自動実行します。")

    add_heading(doc, "9.2　インポート設定項目（csv_import_settings テーブル）", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["設定項目", "説明"],
        [
            ["folder_path", "CSVファイルの格納フォルダ（UNCパス \\\\server\\share\\ も可）"],
            ["filename_pattern", "ファイル名パターン（例：{yyyymm}_売上実績.csv）"],
            ["import_type", "取り込み種別：sales（売上/在庫減算）/ receipt（入荷/在庫加算）/ record_only（記録のみ）"],
            ["col_jan", "JANコードの列番号（0始まり）"],
            ["col_qty", "数量の列番号（0始まり）"],
            ["encoding", "ファイルエンコード：UTF-8-SIG / Shift-JIS / CP932"],
            ["row_filter_dept", "部門コードでフィルタ（指定列の値が一致する行のみ取り込み）"],
            ["row_filter_chain", "チェーンCDでフィルタ"],
            ["monthly_mode", "月次モード：月末の最終ファイルを使用するか毎日取り込むか"],
        ],
        col_widths=[4.5, 12.0]
    )

    add_heading(doc, "9.3　ファイル名パターン仕様", 2)
    add_para(doc, "ファイル名には以下のプレースホルダを使用できます。", kwn=True)
    add_table(doc,
        ["プレースホルダ", "展開例（2026年3月9日の場合）"],
        [
            ["{yyyymmdd}", "20260309"],
            ["{yyyymm}", "202603"],
            ["{yymm}", "2603"],
            ["{yymmdd}", "260309"],
            ["{yyyy}", "2026"],
            ["{mm}", "03"],
            ["{dd}", "09"],
        ],
        col_widths=[5.0, 11.5]
    )

    add_heading(doc, "9.4　処理フロー", 2)
    add_step_table(doc, [
        ("ファイル名解決", "resolve_filename_pattern() でプレースホルダを当日日付に置換してファイルパスを生成。"),
        ("UNC認証", "ネットワーク共有の場合、_unc_server() でサーバー名を取得し認証（net use コマンド）。"),
        ("ファイル読み込み", "指定エンコードでCSVを開き、ヘッダー行をスキップ（skip_rows 設定による）。"),
        ("行フィルタリング", "row_filter_dept / row_filter_chain に一致する行のみを処理対象に絞り込み。"),
        ("重複チェック", "各行を row_hash（SHA256）で識別し、既処理の行はスキップ。"),
        ("JAN正規化", "_normalize_jan() で指数表記・小数点付きJANを変換。"),
        ("数量計算", "import_type が sales の場合：数量 × unit_qty でケース→個換算。"),
        ("在庫更新", "sales → stocks の quantity を減算。receipt → _record_receipt() で在庫加算。"),
        ("集計テーブル更新", "sales_daily_agg に日次集計データをUPSERT（予測キャッシュ無効化）。"),
    ])
    add_note(doc, "重複防止のため、取り込み済みの行ハッシュは csv_import_logs テーブルに保存されます。")
    add_note(doc, "UNCパスへのアクセスにはWindowsの資格情報が必要です（net use での接続認証）。")
    page_break(doc)


def ch10_mail(doc):
    add_heading(doc, "10.　メール通知", 1)

    add_heading(doc, "10.1　メール種別", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["メール種別", "送信タイミング", "送信先"],
        [
            ["発注書メール", "手動発注または自動チェック発注時", "各商品の supplier_email"],
            ["賞味期限アラートメール", "スケジューラー実行時（expiry_alert_days 以内）", "システム管理者メール"],
        ],
        col_widths=[4.0, 6.0, 6.5]
    )

    add_heading(doc, "10.2　発注書メール仕様", 2)
    add_table(doc,
        ["項目", "内容"],
        [
            ["送信方式", "仕入先ごとに1通まとめて送信（queue_order → flush_order_mail）"],
            ["件名", "「【発注書】{仕入先名} {発注日}」形式"],
            ["本文", "商品コード・商品名・発注数量の一覧（テキスト形式）"],
            ["送信結果記録", "order_history.mail_sent（0/1）と mail_result（エラーメッセージ）に保存"],
            ["送信失敗時", "mail_sent=0 で記録。管理画面から再送信可能"],
        ],
        col_widths=[4.0, 12.5]
    )

    add_heading(doc, "10.3　賞味期限アラートメール仕様", 2)
    add_table(doc,
        ["項目", "内容"],
        [
            ["検出条件", "stocks.expiry_date が本日から expiry_alert_days 日以内"],
            ["送信頻度", "スケジューラー実行時に1回（1日1通）"],
            ["通知内容", "アラート対象商品の一覧（JAN・商品名・在庫数量・賞味期限・ロケーション）"],
        ],
        col_widths=[4.0, 12.5]
    )
    page_break(doc)


def ch11_db(doc):
    add_heading(doc, "11.　データベース設計（主要テーブル）", 1)

    add_heading(doc, "11.1　テーブル一覧", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["テーブル名", "用途"],
        [
            ["products", "商品マスタ（発注点・モード・仕入先情報含む）"],
            ["stocks", "在庫台帳（ロット・賞味期限・ロケーション単位）"],
            ["sales_history", "売上履歴（CSV取り込みデータ）"],
            ["sales_daily_agg", "日次売上集計（予測パフォーマンス用キャッシュ）"],
            ["order_history", "発注履歴（手動・自動、メール結果含む）"],
            ["order_pending", "混載保留中の発注（条件待ちキュー）"],
            ["disposed_stocks", "廃棄済み在庫（廃棄理由・ロス金額）"],
            ["stock_movements", "在庫変動ログ（全入出庫・編集・廃棄の監査証跡）"],
            ["stock_transfers", "ロケーション間移動ログ"],
            ["replenishment_history", "補充タスク履歴"],
            ["csv_import_settings", "CSVインポート設定"],
            ["csv_import_logs", "CSVインポート済み行ハッシュ（重複防止）"],
            ["chain_masters", "チェーンマスタ"],
            ["store_masters", "店舗マスタ"],
            ["supplier_code_settings", "チェーン別仕入先コードマッピング"],
            ["product_code_settings", "チェーン別商品コードマッピング"],
            ["promotion_plans", "販促計画（予測の補正に使用）"],
            ["demand_plans", "需要計画（予測の補正に使用）"],
            ["alert_logs", "アラートイベントログ（発注点到達・ロット到達・賞味期限）"],
            ["users", "ユーザーマスタ（認証・権限）"],
        ],
        col_widths=[5.0, 11.5]
    )
    page_break(doc)


def ch12_permissions(doc):
    add_heading(doc, "12.　権限管理", 1)

    add_heading(doc, "12.1　権限チェックの仕組み", 2)
    add_para(doc, "各ルートは Python デコレータで権限を宣言します。セッションに保持された permissions リストと照合します。", kwn=True)
    add_table(doc,
        ["デコレータ", "説明"],
        [
            ["@login_required", "ログイン済みであることを確認。未ログインは /login へリダイレクト"],
            ["@permission_required('xxx')", "指定した権限キーがセッションの permissions に含まれることを確認"],
            ["@admin_required", "admin 権限を持つユーザーのみアクセス可能"],
        ],
        col_widths=[5.5, 11.0]
    )

    add_heading(doc, "12.2　権限とルートの対応", 2)
    add_para(doc, "", kwn=True)
    add_table(doc,
        ["権限キー", "アクセス可能なルートの例"],
        [
            ["dashboard", "GET /dashboard"],
            ["inventory", "GET/POST /inventory, /inventory/edit, /inventory/dispose, /inventory/transfer"],
            ["products", "GET/POST /products, /products/import, /products/export"],
            ["orders", "GET/POST /orders, /orders/send, /orders/auto_check, /backorders"],
            ["forecast", "GET/POST /forecast"],
            ["chains", "GET/POST /chains, /stores"],
            ["admin", "GET/POST /settings, /users（全ルート）"],
        ],
        col_widths=[4.0, 12.5]
    )
    add_note(doc, "権限は users テーブルの permissions カラム（JSON配列）で管理します。")
    add_note(doc, "admin 権限を持つユーザーは全ページにアクセスできます。")


# ═══════════════════════════════════════════════════════════════════════════════
# メイン
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    doc = Document()

    # ページ余白
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    # 既定フォント（游ゴシック）
    doc.styles['Normal'].font.name = '游ゴシック'
    doc.styles['Normal'].font.size = Pt(10)

    build_cover(doc)
    build_toc(doc)
    ch01_overview(doc)
    ch02_auth(doc)
    ch03_dashboard(doc)
    ch04_inventory(doc)
    ch05_products(doc)
    ch06_orders(doc)
    ch07_forecast(doc)
    ch08_chains(doc)
    ch09_csv_import(doc)
    ch10_mail(doc)
    ch11_db(doc)
    ch12_permissions(doc)

    doc.save(OUT)
    print(f"生成完了: {OUT}")

if __name__ == '__main__':
    main()
