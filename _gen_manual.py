# -*- coding: utf-8 -*-
"""在庫管理システム 取扱説明書 生成スクリプト（改ページ完全対策版 v3）"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# =====================================================================
# ヘルパー
# =====================================================================
def set_font(run, size=10.5, bold=False, color=None, font_name="游明朝"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    try:
        run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass

def _kwn(para):
    para.paragraph_format.keep_with_next = True

def _kt(para):
    para.paragraph_format.keep_together = True

def _row_cant_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit')
    cs.set(qn('w:val'), '1')
    trPr.append(cs)

def _row_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), '1')
    trPr.append(th)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _kwn(p)   # 見出しは必ず次要素と同ページ
    _kt(p)
    for run in p.runs:
        run.font.name = "游ゴシック"
        try:
            run._r.rPr.rFonts.set(qn('w:eastAsia'), "游ゴシック")
        except Exception:
            pass
    return p

def add_para(doc, text, indent=False, bold_prefix=None, kwn=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)
    _kt(p)
    if kwn:
        _kwn(p)
    return p

def add_bullet(doc, text, level=0, kwn=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text)
    set_font(r)
    _kt(p)
    if kwn:
        _kwn(p)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """
    テーブルを追加する。
    ★ 呼び出し側で、直前の add_para / add_bullet に kwn=True を付けること。
       見出しは add_heading が自動的に kwn=True にしている。
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダー行：分割禁止 + ページ跨ぎ時に繰り返し
    hdr = table.rows[0]
    _row_cant_split(hdr)
    _row_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _kt(p)
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=(255, 255, 255))

    # データ行：行分割禁止
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        _row_cant_split(row)
        fill = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            _kt(p)
            run = p.add_run(str(val))
            set_font(run, size=9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(1)
    sp.paragraph_format.space_after  = Pt(1)
    return table

def page_break(doc):
    doc.add_page_break()

# =====================================================================
# 表紙
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("在　庫　管　理　シ　ス　テ　ム")
r.font.name = "游ゴシック"; r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(31, 56, 100)
try: r._r.rPr.rFonts.set(qn('w:eastAsia'), "游ゴシック")
except: pass
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("取 扱 説 明 書")
r.font.name = "游ゴシック"; r.font.size = Pt(22); r.font.bold = True
r.font.color.rgb = RGBColor(31, 56, 100)
try: r._r.rPr.rFonts.set(qn('w:eastAsia'), "游ゴシック")
except: pass
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("― 全機能リファレンスガイド ―")
r.font.name = "游明朝"; r.font.size = Pt(14)
try: r._r.rPr.rFonts.set(qn('w:eastAsia'), "游明朝")
except: pass
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 3.0　　2026年4月")
r.font.name = "游明朝"; r.font.size = Pt(12)
try: r._r.rPr.rFonts.set(qn('w:eastAsia'), "游明朝")
except: pass
page_break(doc)

# =====================================================================
# 目次
# =====================================================================
add_heading(doc, "目　次", 1)
toc_items = [
    ("1","はじめに","システム概要・動作環境・ログイン方法"),
    ("2","ダッシュボード","トップ画面・在庫状況サマリー"),
    ("3","商品管理","商品登録・編集・インポート・エクスポート"),
    ("4","発注管理","発注送信・発注履歴・発注残管理・混載"),
    ("5","在庫管理","在庫一覧・入出庫・廃棄・移動・ピッキング・棚補充"),
    ("6","受領管理","入荷受領・受領履歴・インポート"),
    ("7","棚卸管理","棚卸実施・確定・エクスポート"),
    ("8","需要予測・レポート","予測閲覧・発注点自動適用・ABC分析・販促・受注予定"),
    ("9","気温データ管理","気温データ入力・API取得・感応度計算"),
    ("10","52週MD計画","週次販売計画の作成・更新"),
    ("11","チェーン・店舗管理","チェーン・店舗・仕入先CD・商品CD設定"),
    ("12","CSV売上インポート","取込設定・スケジュール・手動実行・ログ確認"),
    ("13","メール設定","受信者管理・テンプレート・テスト送信"),
    ("14","ユーザー管理","ユーザー作成・権限設定・パスワード変更"),
    ("15","システム設定","全般設定・自動更新モード・DBバックアップ"),
    ("16","スケジューラー・自動処理","定時バッチ・自動発注チェック"),
    ("17","権限一覧","ロール・権限の詳細"),
    ("18","用語集","システム内用語の説明"),
]
for num, title, desc in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    _kt(p)
    r = p.add_run(f"第{num}章　{title}")
    set_font(r, bold=True)
    r2 = p.add_run(f"　── {desc}")
    set_font(r2, color=(80, 80, 80))
page_break(doc)

# =====================================================================
# 第1章
# =====================================================================
add_heading(doc, "第1章　はじめに", 1)

add_heading(doc, "1.1　システム概要", 2)
add_para(doc, "本システムは、食品・消耗品などを扱う流通・卸業向けの在庫管理・発注管理システムです。商品マスタ・在庫・発注・売上・需要予測を一元管理し、AI予測や前年実績に基づく発注点自動算出、スケジュール起動による売上CSV自動取込、混載発注管理など、日々の業務を効率化する機能を備えています。")

# 1.2: heading → table （見出しのkwnで表と同ページ）
add_heading(doc, "1.2　動作環境", 2)
add_table(doc,
    ["項目", "内容"],
    [
        ["OS", "Windows 10 / 11（64bit）"],
        ["Webブラウザ", "Google Chrome（最新版）/ Microsoft Edge（最新版）推奨"],
        ["サーバー", "Windows PC上のNSSMサービス（InventorySystem）で常時起動"],
        ["データベース", "PostgreSQL 13以降"],
        ["Python", "Python 3.10〜3.14"],
        ["アクセスURL", "http://localhost:5000（同一PC）または http://[サーバーIP]:5000（LAN内）"],
    ],
    col_widths=[4, 12]
)

add_heading(doc, "1.3　起動・停止方法", 2)
add_para(doc, "【サービスとして自動起動する場合（推奨）】", kwn=True)
add_bullet(doc, "install_service.bat を管理者として実行するとWindowsサービス（InventorySystem）として登録されます。")
add_bullet(doc, "PC起動時に自動的にサーバーが立ち上がります。")
add_para(doc, "【手動で起動する場合】", kwn=True)
add_bullet(doc, "start.bat をダブルクリックして起動します。コンソール画面が表示されている間サーバーが動作します。")
add_bullet(doc, "停止するにはコンソール画面を閉じるか Ctrl+C を押します。")
add_para(doc, "【サービスの再起動（設定変更後など）】", kwn=True)
add_bullet(doc, r"管理者PowerShellで nssm\nssm.exe stop InventorySystem → nssm\nssm.exe start InventorySystem を実行します。")

# 1.4: heading → para → table
# ★ paraにkwn=Trueを付けて見出し→段落→表を同ページに固定
add_heading(doc, "1.4　ログイン方法", 2)
add_para(doc, "ブラウザで http://localhost:5000 にアクセスするとログイン画面が表示されます。", kwn=True)
add_table(doc,
    ["項目", "説明"],
    [
        ["ユーザー名", "管理者から発行されたユーザー名を入力します"],
        ["パスワード", "パスワードを入力します（初回ログイン後は変更が必要な場合があります）"],
        ["ログイン失敗制限", "1分間に10回まで（超過するとブロック）"],
    ],
    col_widths=[4, 12]
)
add_para(doc, "※ 初回ログイン時にパスワード変更を求められる場合があります。8文字以上で設定してください。")
page_break(doc)

# =====================================================================
# 第2章
# =====================================================================
add_heading(doc, "第2章　ダッシュボード", 1)
add_para(doc, "ログイン後のトップ画面です。在庫状況の概要をひと目で確認できます。")

add_heading(doc, "2.1　表示内容", 2)
add_table(doc,
    ["表示エリア", "内容"],
    [
        ["総商品数", "有効な商品の総件数"],
        ["発注検討商品数", "在庫が発注点以下になっている商品の件数"],
        ["期限アラート件数", "賞味期限切れ間近の在庫件数"],
        ["発注検討商品リスト", "JAN・商品名・在庫数・発注点を一覧表示。発注画面へのリンク付き"],
        ["最近の発注履歴", "直近10件の発注記録（発注日・仕入先・商品・数量）"],
        ["最近のアラート", "直近6件のアラート（期限超過・在庫切れなど）"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "2.2　操作", 2)
add_bullet(doc, "「発注する」ボタン：発注検討商品の一覧から発注管理画面へ移動できます。")
add_bullet(doc, "各商品名のリンク：商品詳細・編集画面に移動します。")
page_break(doc)

# =====================================================================
# 第3章
# =====================================================================
add_heading(doc, "第3章　商品管理", 1)
add_para(doc, "商品マスタの登録・編集・検索・インポート・エクスポートを行います。")

# 3.1: heading → para(kwn) → para(kwn) → table
add_heading(doc, "3.1　商品一覧", 2)
add_para(doc, "URL: /products　　権限: products", kwn=True)
add_para(doc, "有効商品の一覧を表示します。検索ボックスでJAN・商品コード・商品名・仕入先名を絞り込めます。", kwn=True)
add_table(doc,
    ["表示列", "説明"],
    [
        ["JAN", "JANコード（バーコード番号）"],
        ["商品CD", "社内商品コード"],
        ["商品名", "商品の名称"],
        ["仕入先CD / 名", "仕入先コードと名称"],
        ["在庫数", "現在の在庫合計（stocks テーブルより集計）"],
        ["発注点", "この数量以下になると発注検討対象になる"],
        ["発注数量", "1回の発注で送る数量"],
        ["発注点自動更新", "手動 / AIモード / 前年実績 の3種類"],
        ["入数", "1ケース（箱）に入る個数"],
        ["原価 / 売価", "仕入単価 / 販売単価"],
    ],
    col_widths=[4, 12]
)

# 3.2: heading → para(kwn) → para(kwn) → table
add_heading(doc, "3.2　商品新規登録", 2)
add_para(doc, "URL: /products/new　　権限: admin", kwn=True)
add_para(doc, "「新規登録」ボタンをクリックして入力フォームを開きます。", kwn=True)
add_table(doc,
    ["入力項目", "必須", "説明"],
    [
        ["仕入先CD", "○", "仕入先の識別コード"],
        ["仕入先名", "○", "仕入先の名称"],
        ["仕入先メール", "", "発注メール送信先（複数はカンマ区切り）"],
        ["JANコード", "○", "バーコード番号（システム内で一意）"],
        ["商品コード", "○", "社内管理コード"],
        ["商品名", "○", "商品の名称"],
        ["入数（unit_qty）", "", "1ケースに入る個数（デフォルト1）"],
        ["発注単位", "", "発注時の最小ロット数（デフォルト1）"],
        ["発注数量", "", "通常の1回発注数量"],
        ["発注点", "", "在庫がこの数以下で発注検討に上がる"],
        ["発注点自動更新", "", "0=手動 / 1=AIモード / 2=前年実績"],
        ["リードタイム（日）", "", "発注から入荷までの日数"],
        ["安全係数", "", "需要変動のバッファ係数（例: 1.5）"],
        ["ロット数", "", "発注ロット単位"],
        ["賞味期限日数", "", "製造日からの賞味期限日数"],
        ["期限アラート日数", "", "期限まで残りこの日数でアラート発生"],
        ["混載グループ名", "", "混載発注をまとめるグループID"],
        ["混載ロットモード", "", "gte（以上）/ unit（倍数）"],
        ["混載ケース数", "", "混載発注の最低ケース数"],
        ["強制発注日数", "", "混載条件未達でもこの日数経過で強制発注"],
        ["原価（仕入単価）", "○", "商品の仕入れ価格"],
        ["売価", "○", "商品の販売価格"],
        ["ロケーションコード", "", "倉庫内の保管場所コード"],
        ["棚面数", "", "棚に並べる面数"],
        ["棚補充点", "", "棚在庫がこれ以下で補充指示を出す"],
        ["手動調整係数", "", "予測値に乗算する手動補正値（例: 1.2）"],
    ],
    col_widths=[4.5, 1.5, 10]
)

add_heading(doc, "3.3　商品編集・無効化", 2)
add_para(doc, "URL: /products/<id>/edit　　権限: admin")
add_bullet(doc, "商品一覧の「編集」ボタンから各商品の情報を変更できます。")
add_bullet(doc, "「無効化」ボタンを押すと商品が非表示になります（データは残ります）。")
add_bullet(doc, "無効化された商品は「無効商品を表示」チェックで確認・復活できます。")

# 3.4: heading → para(kwn) → para(kwn) → table
add_heading(doc, "3.4　商品インポート", 2)
add_para(doc, "URL: /products/import　　権限: admin", kwn=True)
add_para(doc, "ExcelまたはCSVファイルで商品を一括登録・更新できます。", kwn=True)
add_table(doc,
    ["設定項目", "説明"],
    [
        ["ファイル形式", "CSV（UTF-8/SJIS）またはXLSX/XLS"],
        ["インポートモード", "upsert: 既存商品は更新、なければ新規追加\nadd_only: 新規追加のみ（既存はスキップ）"],
        ["キー照合順序", "JANコードで照合 → なければ商品CDで照合"],
        ["発注点自動更新値", "「0 手動」「1 AIモード」「2 前年実績」等のテキストも自動変換"],
    ],
    col_widths=[4.5, 11.5]
)
add_para(doc, "テンプレートファイルは「テンプレートDL」ボタンからダウンロードできます（凡例シート付き）。")

add_heading(doc, "3.5　商品エクスポート", 2)
add_para(doc, "URL: /products/export　　権限: admin")
add_bullet(doc, "「CSV出力」または「Excel出力」ボタンで全有効商品データをダウンロードできます。")
add_bullet(doc, "エクスポートしたファイルはそのままインポート用テンプレートとして使用できます。")
page_break(doc)

# =====================================================================
# 第4章
# =====================================================================
add_heading(doc, "第4章　発注管理", 1)
add_para(doc, "仕入先への発注送信・履歴確認・部分入荷管理・混載発注管理を行います。")

# 4.1: heading → para(kwn) → table
add_heading(doc, "4.1　発注メイン画面", 2)
add_para(doc, "URL: /orders　　権限: orders", kwn=True)
add_table(doc,
    ["タブ", "内容"],
    [
        ["発注検討商品", "在庫が発注点以下の商品一覧。チェックを入れて「発注する」で発注送信"],
        ["発注済み商品", "既に発注済み（ordered_at が設定されている）の商品一覧"],
        ["全商品", "全商品の在庫・発注状況を一覧表示"],
    ],
    col_widths=[4, 12]
)

add_heading(doc, "4.2　発注する", 2)
add_para(doc, "発注検討商品タブでチェックを入れて「発注する」ボタンを押します。")
add_bullet(doc, "数量はフォームで変更可能です（デフォルトは商品マスタの発注数量）。")
add_bullet(doc, "発注内容は order_history テーブルに記録されます。")
add_bullet(doc, "設定されたメールアドレスに発注メールが送信されます。")
add_bullet(doc, "発注後、商品の「発注済み」フラグが立ち、発注済みタブに移動します。")
add_bullet(doc, "入数（unit_qty）で自動的にケース単位に切り上げられます。")

add_heading(doc, "4.3　自動チェック実行", 2)
add_para(doc, "「自動チェック実行」ボタンを押すと在庫状況を自動チェックします。")
add_bullet(doc, "発注点以下の商品を自動検出します。")
add_bullet(doc, "混載グループ設定がある商品は order_pending（保留）テーブルに追加されます。")
add_bullet(doc, "強制発注日数を超えた保留商品は自動的に発注処理されます。")

# 4.4: heading → para(kwn) → table
add_heading(doc, "4.4　発注履歴", 2)
add_para(doc, "URL: /order_history　　権限: order_history", kwn=True)
add_table(doc,
    ["表示列", "説明"],
    [
        ["発注日時", "発注を実行した日時"],
        ["仕入先", "発注先の仕入先名"],
        ["商品", "発注商品名とJAN"],
        ["発注数量", "発注した数量"],
        ["トリガー種別", "manual（手動）/ reorder（自動）/ forecast（予測）/ forced（強制）"],
        ["メール送信状態", "sent（送信済）/ failed（失敗）/ none（未送信）"],
        ["期待入荷日", "発注日 + リードタイム日数"],
    ],
    col_widths=[4, 12]
)
add_para(doc, "検索・フィルタ・ページング（100件/ページ）が使用できます。")

# 4.5: heading → para(kwn) → para(kwn) → table
add_heading(doc, "4.5　発注残・部分入荷管理", 2)
add_para(doc, "URL: /orders/backorders　　権限: orders", kwn=True)
add_para(doc, "発注済みだが入荷が完了していない商品を管理します。", kwn=True)
add_table(doc,
    ["ステータス", "条件"],
    [
        ["完了", "全量入荷済み"],
        ["入荷接近", "期待入荷日まで3日以内"],
        ["遅延", "期待入荷日を超過している"],
        ["長期未着", "発注から60日以上経過"],
        ["正常", "上記以外"],
    ],
    col_widths=[4, 12]
)
add_para(doc, "【部分受領登録】", kwn=True)
add_bullet(doc, "「受領登録」ボタンで入荷数量・賞味期限・ロット番号・保管場所を入力して在庫に加算します。")
add_bullet(doc, "残数量を超えた入力はできません。全量入荷後は発注済みフラグが自動解除されます。")
add_para(doc, "【未納クローズ】", kwn=True)
add_bullet(doc, "「未納クローズ」ボタンで発注残を終了扱いにします（理由を入力）。")

# 4.6: heading → para(kwn) → table
add_heading(doc, "4.6　混載発注管理", 2)
add_para(doc, "URL: /orders/pending_force_group_edit　　権限: orders", kwn=True)
add_table(doc,
    ["混載ロットモード", "説明"],
    [
        ["gte（以上）", "グループ合計ケース数が mixed_lot_cases 以上になったら発注"],
        ["unit（倍数）", "グループ合計ケース数が mixed_lot_cases の倍数になったら発注"],
    ],
    col_widths=[4, 12]
)
add_bullet(doc, "強制発注日数（mixed_force_days）を超えた場合は条件に関わらず強制発注されます。")
add_bullet(doc, "「グループ一括発注」で同一グループをまとめて発注できます。")
add_bullet(doc, "「手動調整発注」で各商品の数量を調整してから発注できます。")
page_break(doc)

# =====================================================================
# 第5章
# =====================================================================
add_heading(doc, "第5章　在庫管理", 1)
add_para(doc, "在庫の参照・直接編集・廃棄・ロケーション移動・ピッキング・棚補充を行います。")

# 5.1: heading → para(kwn) → table
add_heading(doc, "5.1　在庫一覧", 2)
add_para(doc, "URL: /inventory　　権限: inventory", kwn=True)
add_table(doc,
    ["表示列", "説明"],
    [
        ["JAN / 商品名", "商品の識別情報"],
        ["仕入先", "仕入先名"],
        ["在庫数", "現在の数量"],
        ["賞味期限", "期限切れ間近（30日以内）は橙色で表示"],
        ["ロット番号", "入荷時のロット識別番号"],
        ["保管場所（ロケーション）", "倉庫内の保管場所コード"],
        ["発注点", "在庫が発注点以下で発注検討に上がる"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "5.2　在庫直接編集", 2)
add_para(doc, "「編集」ボタンで在庫数量・賞味期限・ロット番号・ロケーションを直接変更できます。")
add_bullet(doc, "変更前後の数量差は stock_movements テーブルに調整記録として自動保存されます。")

# 5.3: heading → para(kwn) → table
add_heading(doc, "5.3　在庫廃棄・退避", 2)
add_para(doc, "「廃棄・退避」ボタンで在庫を廃棄扱いにします。", kwn=True)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["廃棄理由種別", "期限切れ / 破損 / 品質不良 / 返品 / 退避 / その他"],
        ["廃棄理由詳細", "自由記述"],
        ["廃棄数量", "廃棄する数量"],
    ],
    col_widths=[5, 11]
)
add_bullet(doc, "廃棄後は disposed_stocks テーブルに記録されます。原価×廃棄数量でロス金額が自動計算されます。")
add_bullet(doc, "廃棄を取り消す場合は「復元」ボタンを使います（/inventory/disposed で確認）。")

add_heading(doc, "5.4　廃棄在庫一覧", 2)
add_para(doc, "URL: /inventory/disposed　　権限: ログイン済")
add_bullet(doc, "廃棄日・商品・数量・ロス金額・廃棄理由を確認できます。CSV/Excelでエクスポートできます。")
add_bullet(doc, "「復元」ボタンで在庫に戻すことができます。")

# 5.5: heading → para(kwn) → table
add_heading(doc, "5.5　ロケーション移動", 2)
add_para(doc, "URL: /inventory/transfers　　権限: inventory", kwn=True)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["移動元在庫", "移動する在庫を選択（JAN・期限・ロット・現在の場所が表示される）"],
        ["移動先ロケーション", "移動後の保管場所コード"],
        ["移動数量", "移動する数量（元在庫数を超えられない）"],
        ["備考", "移動理由など"],
    ],
    col_widths=[5, 11]
)
add_bullet(doc, "移動履歴は過去300件表示されます。stock_movements に transfer_out / transfer_in として記録されます。")

add_heading(doc, "5.6　ピッキング計画", 2)
add_para(doc, "URL: /inventory/picking　　権限: inventory")
add_bullet(doc, "仕入先別・期間別でフィルタして必要数量・ロケーション・賞味期限を確認できます。")
add_bullet(doc, "「CSV出力」「Excel出力」「印刷」ボタンでピッキングリストを出力できます。")

# 5.7: heading → para(kwn) → table
add_heading(doc, "5.7　棚補充管理", 2)
add_para(doc, "URL: /inventory/replenishment　　権限: inventory", kwn=True)
add_table(doc,
    ["ステータス", "説明"],
    [
        ["planned", "補充指示が作成された"],
        ["partial", "一部補充済み"],
        ["done", "補充完了"],
    ],
    col_widths=[4, 12]
)
add_bullet(doc, "「補充指示作成」で補充数量・補充元ロケーション・補充先棚を指定します。")
add_bullet(doc, "「補充完了登録」で実際に補充した数量を記録します。")
page_break(doc)

# =====================================================================
# 第6章
# =====================================================================
add_heading(doc, "第6章　受領管理", 1)
add_para(doc, "仕入先からの入荷商品を在庫に登録します。")

# 6.1: heading → para(kwn) → table
add_heading(doc, "6.1　受領入力", 2)
add_para(doc, "URL: /receipt　　権限: receipt", kwn=True)
add_table(doc,
    ["入力項目", "必須", "説明"],
    [
        ["JAN / 商品CD", "○", "受領する商品のJANまたは商品コード"],
        ["受領数量", "○", "入荷した数量"],
        ["賞味期限", "○", "受領した商品の賞味期限（YYYY-MM-DD）"],
        ["ロット番号", "", "入荷ロットの識別番号"],
        ["保管場所", "", "倉庫内の格納場所コード"],
        ["メール通知", "", "受領メールを送信する受信者を選択"],
    ],
    col_widths=[4, 1.5, 10.5]
)
add_bullet(doc, "受領後、stocks テーブルに在庫が追加され、stock_movements に receipt として記録されます。")

add_heading(doc, "6.2　受領履歴", 2)
add_para(doc, "URL: /receipt/history　　権限: order_history")
add_bullet(doc, "受領日・商品・数量・メール送信状態を確認できます。")
add_bullet(doc, "「削除」ボタンで在庫を減算して取り消しができます。")

add_heading(doc, "6.3　受領インポート", 2)
add_para(doc, "URL: /receipt/import　　権限: admin")
add_bullet(doc, "CSV/Excelファイルでまとめて受領登録できます。テンプレートは「テンプレートDL」からDLできます。")
add_bullet(doc, "対応列：JAN、商品名、仕入先CD/名、数量、賞味期限、ロット番号、ロケーション")
page_break(doc)

# =====================================================================
# 第7章
# =====================================================================
add_heading(doc, "第7章　棚卸管理", 1)
add_para(doc, "実際に数えた在庫数をシステムに記録し、差異を管理します。")

# 7.1: heading → para(kwn) → table
add_heading(doc, "7.1　棚卸実施", 2)
add_para(doc, "URL: /stocktake　　権限: stocktake", kwn=True)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["棚卸日付", "棚卸を実施した日付"],
        ["商品（JAN）", "棚卸対象の商品"],
        ["実棚在庫数", "実際に数えた数量"],
        ["差分理由（カテゴリ）", "盗難 / 破損 / 計上誤り / 入出庫漏れ / その他"],
        ["差分理由（詳細）", "自由記述"],
        ["期限別内訳", "複数ロット・期限がある場合の詳細情報"],
    ],
    col_widths=[5, 11]
)
add_para(doc, "システム在庫と実棚在庫の差分（diff_qty）が自動計算されます。")

add_heading(doc, "7.2　棚卸確定", 2)
add_bullet(doc, "「棚卸確定」を押すと stocks テーブルの数量が実棚数に更新されます。確定後は元に戻せません。")
add_bullet(doc, "stock_movements テーブルに adjust として調整記録が保存されます。")

add_heading(doc, "7.3　棚卸エクスポート", 2)
add_bullet(doc, "「CSV出力」「Excel出力」ボタンで棚卸結果をダウンロードできます。")
add_bullet(doc, "出力列：棚卸日・仕入先・商品・システム在庫・実棚在庫・差分・理由")
page_break(doc)

# =====================================================================
# 第8章
# =====================================================================
add_heading(doc, "第8章　需要予測・レポート", 1)
add_para(doc, "売上履歴をもとに将来の需要を予測し、最適な発注点・発注数量を提案します。")

# 8.1: heading → para(kwn) → table
add_heading(doc, "8.1　需要予測メイン画面", 2)
add_para(doc, "URL: /reports/forecast　　権限: forecast", kwn=True)
add_table(doc,
    ["表示列", "説明"],
    [
        ["ABC分析ランク", "売上構成比でA（上位70%）/ B（71〜90%）/ C（91%〜）に分類"],
        ["Q25 / Q50 / Q75 日次予測", "日別販売数の25/50/75パーセンタイル予測値"],
        ["動的安全在庫（DSS）", "需要の変動幅に応じた推奨安全在庫量"],
        ["提案発注点", "Q75日次予測 × リードタイム + DSS で算出"],
        ["提案発注数量", "Q50予測 × リードタイム で算出"],
        ["計算アルゴリズム", "SMA（単純移動平均）/ ARIMA / ETS など"],
    ],
    col_widths=[5, 11]
)

# 8.2: heading → para(kwn) → table
add_heading(doc, "8.2　発注点自動適用", 2)
add_para(doc, "「一括反映」ボタンで予測した発注点・発注数量を商品マスタに書き込みます。", kwn=True)
add_table(doc,
    ["項目", "説明"],
    [
        ["対象モード", "AIモード（reorder_auto=1）または前年実績モード（reorder_auto=2）の商品のみ適用"],
        ["計算モード（AI）", "予測エンジン（SMA/ARIMA/ETS）で計算した発注点を適用"],
        ["計算モード（前年実績）", "前年同月の販売実績をもとに計算した発注点を適用"],
        ["処理状況確認", "バックグラウンド処理のため進捗バーで確認可能"],
    ],
    col_widths=[5, 11]
)
add_para(doc, "※ reorder_auto=0（手動）の商品には適用されません。")

# 8.3: heading → para(kwn) → table
add_heading(doc, "8.3　ABC分析", 2)
add_para(doc, "URL: /reports/abc　　権限: forecast", kwn=True)
add_table(doc,
    ["ランク", "デフォルト基準", "説明"],
    [
        ["A", "上位70%", "売上高の高い主力商品。欠品防止を最優先"],
        ["B", "71〜90%", "中間的な重要度の商品"],
        ["C", "91%〜", "売上貢献の少ない商品。在庫圧縮対象"],
    ],
    col_widths=[3, 4, 9]
)
add_bullet(doc, "「Excelエクスポート」ボタンでABC分析結果をダウンロードできます。")

# 8.4: heading → para(kwn) → para(kwn) → table
add_heading(doc, "8.4　販促計画登録", 2)
add_para(doc, "URL: /reports/forecast/promotions　　権限: forecast", kwn=True)
add_para(doc, "セール・キャンペーン等の販促イベントを登録し、需要予測に反映します。", kwn=True)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["JAN", "販促対象商品のJANコード"],
        ["販促日", "販促実施日（YYYY-MM-DD）"],
        ["販促名", "キャンペーン名など"],
        ["売上上昇係数（uplift_factor）", "通常売上に対する倍率（例: 1.5 = 150%に増加）"],
    ],
    col_widths=[5.5, 10.5]
)
add_bullet(doc, "テンプレートXLSXをダウンロードして一括インポートすることもできます。")

# 8.5: heading → para(kwn) → para(kwn) → table
add_heading(doc, "8.5　受注予定登録", 2)
add_para(doc, "URL: /reports/forecast/demands　　権限: forecast", kwn=True)
add_para(doc, "確定している大口注文やイベント需要を事前登録し、予測に加算します。", kwn=True)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["JAN", "対象商品のJANコード"],
        ["需要日", "需要が発生する日付"],
        ["需要数量", "予定数量"],
        ["需要種別", "order（受注）/ promo（販促）/ custom（その他）"],
        ["顧客名", "発注元の顧客名（任意）"],
    ],
    col_widths=[5.5, 10.5]
)
add_bullet(doc, "テンプレートXLSXをダウンロードして一括インポートすることもできます。")
add_bullet(doc, "個別削除または「全件削除」で一括クリアできます。")

add_heading(doc, "8.6　問屋向け予測", 2)
add_para(doc, "URL: /reports/forecast/wholesale　　権限: forecast")
add_bullet(doc, "気温感応度（商品別の気温との相関係数）を加味した発注点を表示します。")
add_bullet(doc, "「商品マスタに反映」ボタンで気温調整後の発注点を一括適用できます。")
page_break(doc)

# =====================================================================
# 第9章
# =====================================================================
add_heading(doc, "第9章　気温データ管理", 1)
add_para(doc, "URL: /reports/weather　　権限: forecast")
add_para(doc, "気温と売上の相関を分析し、季節性・気温変動を需要予測に活かします。")

# 9.1: heading → table （見出しのkwnで表と同ページ）
add_heading(doc, "9.1　気温データ入力", 2)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["観測日", "気温を記録した日付"],
        ["ロケーション", "観測地点（例: 東京、大阪、宮崎）"],
        ["平均気温（℃）", "その日の平均気温"],
        ["最高気温（℃）", "その日の最高気温"],
        ["最低気温（℃）", "その日の最低気温"],
        ["降水量（mm）", "その日の降水量"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "9.2　気象APIによる自動取得", 2)
add_bullet(doc, "「API取得」ボタンで気象APIから自動的に気温データを取得・保存します。")
add_bullet(doc, "スケジューラーにより毎日03:00に自動取得されます（設定で時刻変更可）。")

add_heading(doc, "9.3　気温感応度再計算", 2)
add_bullet(doc, "「感応度再計算」ボタンで商品ごとに気温との相関係数（r²）を再計算します。")
add_bullet(doc, "相関係数が高い商品は気温変動の影響を受けやすい商品です。")

add_heading(doc, "9.4　CSVインポート", 2)
add_bullet(doc, "「CSVインポート」で気象庁等からダウンロードしたCSVを一括取込できます。")
page_break(doc)

# =====================================================================
# 第10章
# =====================================================================
add_heading(doc, "第10章　52週MD計画", 1)
add_para(doc, "URL: /reports/weekly_md　　権限: forecast")
add_para(doc, "1年間（52週）の販売計画（MD計画）を週単位で管理します。")

# 10.1: heading → table
add_heading(doc, "10.1　計画生成・更新", 2)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["会計年度（fiscal_year）", "対象の会計年度（例: 2026）"],
        ["週番号（week_no）", "第1週〜第52週"],
        ["計画数量（plan_qty）", "週の販売計画数量"],
        ["実績数量（actual_qty）", "実際の販売数量（事後入力）"],
    ],
    col_widths=[5.5, 10.5]
)
add_bullet(doc, "「計画生成」ボタンで指定年度の52週分の雛形を一括作成します。")
add_bullet(doc, "「Excelエクスポート」で計画シートをダウンロードできます。")
page_break(doc)

# =====================================================================
# 第11章
# =====================================================================
add_heading(doc, "第11章　チェーン・店舗管理", 1)
add_para(doc, "URL: /chains　　権限: chains")
add_para(doc, "取引先チェーン・店舗のマスタと在庫引き当て除外設定を管理します。")

add_heading(doc, "11.1　チェーンマスタ", 2)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["チェーンCD", "チェーン識別コード"],
        ["チェーン名", "チェーン店名"],
        ["在庫引き当て除外", "このチェーンへの出荷を在庫引き当て計算から除外するか"],
    ],
    col_widths=[5, 11]
)
add_bullet(doc, "テンプレートDLおよびXLSXインポートで一括登録できます。")

add_heading(doc, "11.2　店舗マスタ", 2)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["店舗CD", "店舗識別コード"],
        ["店舗名", "店舗名"],
        ["チェーンCD", "所属するチェーンのコード"],
        ["取引先名", "取引先の正式名"],
        ["在庫引き当て除外", "この店舗への出荷を引き当て除外するか"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "11.3　仕入先CD別設定", 2)
add_para(doc, "特定のチェーン・店舗に対して、特定の仕入先CDの在庫を引き当て除外に設定します。テンプレートDLおよびXLSXインポートで一括登録できます。")

add_heading(doc, "11.4　商品CD別設定", 2)
add_para(doc, "特定のチェーン・店舗に対して、特定の商品CD（またはJAN）の在庫を引き当て除外に設定します。テンプレートDLおよびXLSXインポートで一括登録できます。")
page_break(doc)

# =====================================================================
# 第12章
# =====================================================================
add_heading(doc, "第12章　CSV売上インポート設定", 1)
add_para(doc, "URL: /csv　　権限: admin")
add_para(doc, "POSシステム・受注システム等からの売上CSVを自動的に取り込む設定を管理します。")

# 12.1: heading → para(kwn) → table
add_heading(doc, "12.1　インポート設定の作成", 2)
add_para(doc, "URL: /csv/new", kwn=True)
add_table(doc,
    ["設定項目", "説明"],
    [
        ["設定名", "この設定の名前（例: POSシステム取込）"],
        ["インポート種別", "sales（売上）/ receipt（受領）/ その他"],
        ["フォルダパス", "取込元CSVが置かれるフォルダのパス"],
        ["ネットユーザー/パスワード", "ネットワーク共有フォルダへのアクセス認証情報"],
        ["ファイル名パターン", "ファイル名のパターン（{yyyymmdd} で日付部分を置換）\n例: sales_{yyyymmdd}.csv"],
        ["文字エンコード", "UTF-8 / Shift-JIS / EUC-JP など"],
        ["スケジュール実行時刻", "自動実行する時刻（HH:MM形式、複数設定可）"],
        ["CSVカラムマッピング", "各列が何のデータか（JAN/数量/日付/チェーンCD/店舗CDなど）"],
        ["フィルタCD", "特定の担当者コード等でフィルタリング"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "12.2　手動実行", 2)
add_bullet(doc, "「実行」ボタンで設定したCSVを即時取り込みます。")
add_bullet(doc, "「全設定一括実行」ですべての設定を一括処理します。")
add_bullet(doc, "「月末集計実行」で月次集計処理を実行します。")
add_bullet(doc, "「再インポート」で指定期間のデータを再読み込みします。")

# 12.3: heading → para(kwn) → table
add_heading(doc, "12.3　インポートログ確認", 2)
add_para(doc, "URL: /csv/import_detail", kwn=True)
add_table(doc,
    ["表示項目", "説明"],
    [
        ["ファイル名", "処理したCSVファイル名"],
        ["OK行数", "正常に取り込めた行数"],
        ["エラー行数", "エラーが発生した行数"],
        ["ステータス", "success / partial / error / running"],
        ["実行日時", "処理を実行した日時"],
    ],
    col_widths=[5, 11]
)
add_bullet(doc, "エラー行の詳細（行番号・エラー内容）を確認できます。")
add_bullet(doc, "ログ保持期間は設定画面で変更できます（デフォルト: 3ヶ月）。")
page_break(doc)

# =====================================================================
# 第13章
# =====================================================================
add_heading(doc, "第13章　メール設定", 1)
add_para(doc, "発注メール・期限アラートメールの受信者とテンプレートを管理します。")

# 13.1: heading → para(kwn) → table
add_heading(doc, "13.1　メール受信者管理", 2)
add_para(doc, "URL: /recipients　　権限: admin", kwn=True)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["名前", "受信者の氏名"],
        ["メールアドレス", "送信先メールアドレス"],
        ["送信タイプ", "order（発注メールのみ）/ expiry（期限アラートのみ）/ both（両方）"],
        ["仕入先CD", "指定した場合、その仕入先の発注メールのみ受信"],
        ["有効/無効", "無効にするとメールが送信されなくなる"],
    ],
    col_widths=[4.5, 11.5]
)
add_bullet(doc, "テンプレートDLおよびXLSXインポートで一括登録できます。")

# 13.2: heading → para(kwn) → table
add_heading(doc, "13.2　メールテンプレート", 2)
add_para(doc, "URL: /settings/mail_templates　　権限: admin", kwn=True)
add_table(doc,
    ["設定項目", "説明"],
    [
        ["メール種別", "order（発注）/ expiry（期限アラート）"],
        ["件名（subject）", "メールの件名。{date}等のプレースホルダが使用可能"],
        ["本文ヘッダー", "メール本文の冒頭部分"],
        ["本文明細行テンプレート", "商品1行分のフォーマット（{product_name}等を使用）"],
        ["本文フッター", "メール本文の末尾部分"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "13.3　テストメール送信", 2)
add_bullet(doc, "「テスト送信」ボタンで設定したテンプレートのテストメールを送信して確認できます。")
add_bullet(doc, "SMTP設定（ホスト・ポート・アカウント）は /settings（システム設定）で行います。")
page_break(doc)

# =====================================================================
# 第14章
# =====================================================================
add_heading(doc, "第14章　ユーザー管理", 1)
add_para(doc, "URL: /users　　権限: admin")
add_para(doc, "システムにアクセスするユーザーのアカウント・権限を管理します。")

# 14.1: heading → table
add_heading(doc, "14.1　ユーザー新規作成", 2)
add_table(doc,
    ["入力項目", "説明"],
    [
        ["ユーザー名", "ログインに使用するID（システム内で一意）"],
        ["パスワード", "8文字以上で設定。SHA256でハッシュ化して保存"],
        ["ロール", "admin（管理者）/ user（一般ユーザー）"],
    ],
    col_widths=[5, 11]
)

# 14.2: heading → para(kwn) → table
add_heading(doc, "14.2　権限設定", 2)
add_para(doc, "各ユーザーに対して機能ごとのアクセス権限を設定します。", kwn=True)
add_table(doc,
    ["権限キー", "アクセスできる機能"],
    [
        ["dashboard", "ダッシュボード"],
        ["products", "商品管理（一覧・検索）。登録・編集・削除は admin のみ"],
        ["orders", "発注管理・発注操作"],
        ["inventory", "在庫管理・廃棄・移動"],
        ["forecast", "需要予測・レポート"],
        ["chains", "チェーン・店舗管理"],
        ["order_history", "発注履歴・受領管理"],
        ["reports", "各種レポート"],
        ["receipt", "受領入力"],
        ["csv", "CSV設定・インポート"],
        ["stocktake", "棚卸管理"],
        ["users", "ユーザー管理（admin権限と同等）"],
        ["settings", "システム設定"],
    ],
    col_widths=[4, 12]
)
add_para(doc, "※ admin ロールは自動的にすべての権限を持ちます。")

add_heading(doc, "14.3　ユーザー操作", 2)
add_bullet(doc, "「有効/無効切替」ボタンでアカウントを一時停止・再開できます。")
add_bullet(doc, "「パスワード変更」ボタンで管理者が任意ユーザーのパスワードを変更できます。")
add_bullet(doc, "「削除」ボタンでユーザーを完全削除します（元に戻せません）。")
page_break(doc)

# =====================================================================
# 第15章
# =====================================================================
add_heading(doc, "第15章　システム設定", 1)
add_para(doc, "URL: /settings　　権限: admin")
add_para(doc, "システム全体の動作に関わるパラメータを設定します。")

# 15.1: heading → table
add_heading(doc, "15.1　全般設定", 2)
add_table(doc,
    ["設定項目", "説明"],
    [
        ["会社名", "画面タイトル・メールに表示される会社名"],
        ["ロゴ画像", "ナビゲーションに表示するロゴ（PNG/JPG）"],
        ["発注点自動更新モード", "manual / ai / ly\n※ 設定保存時に全商品の reorder_auto が一括更新される"],
        ["予測アルゴリズム", "SMA（単純移動平均）/ ARIMA / ETS から選択"],
        ["AI予測モード", "有効にするとAIアルゴリズムが利用可能"],
        ["手動調整係数", "有効にすると manual_adjust_factor が予測に乗算される"],
        ["分位点発注点モード", "有効にすると Q75 ベースで発注点を算出（リスク重視）"],
        ["安全係数（Z値）", "需要変動バッファの倍率。1.65（信頼度95%）が標準"],
        ["ABC閾値A / B", "A/Bランクを決める売上構成比（デフォルト: A=70%, B=90%）"],
        ["気象ロケーション", "気温データ取得の観測地点（例: 宮崎）"],
        ["在庫保持月数", "sales_historyの保持期間。古いデータは自動削除"],
        ["廃棄ログ保持月数", "廃棄記録（disposed_stocks）の保持期間"],
    ],
    col_widths=[5.5, 10.5]
)

# 15.2: heading → table
add_heading(doc, "15.2　SMTP設定（メール送信）", 2)
add_table(doc,
    ["設定項目", "説明"],
    [
        ["SMTPホスト", "メールサーバーのホスト名またはIPアドレス"],
        ["SMTPポート", "接続ポート（一般的に 25 / 587 / 465）"],
        ["SMTPユーザー", "認証ユーザー名"],
        ["SMTPパスワード", "認証パスワード"],
        ["送信者アドレス", "Fromに表示されるメールアドレス"],
    ],
    col_widths=[5.5, 10.5]
)

# 15.3: heading → para(kwn) → table
add_heading(doc, "15.3　発注点自動更新モードの一括切替", 2)
add_para(doc, "設定画面の「発注点自動更新モード」を変更して「設定保存」を押すと全商品のモードが一括更新されます。", kwn=True)
add_table(doc,
    ["モード", "値", "動作"],
    [
        ["手動（manual）", "0", "自動更新しない。発注点・数量は手動で設定した値を使い続ける"],
        ["AIモード（ai）", "1", "毎月1日にAI予測エンジンで発注点・数量を自動更新"],
        ["前年実績（ly）", "2", "毎月1日に前年同月の販売実績をもとに発注点・数量を自動更新"],
    ],
    col_widths=[3.5, 2, 10.5]
)

add_heading(doc, "15.4　バックアップ・復元", 2)
add_bullet(doc, "「バックアップ」ボタンで全テーブルのデータをSQLファイルとしてダウンロードできます。")
add_bullet(doc, "「復元」ではSQLファイルをアップロードしてデータを復元します（既存データを上書き）。")

add_heading(doc, "15.5　一括削除（古いデータ整理）", 2)
add_bullet(doc, "古い発注履歴・棚卸ログ・エラーログ・インポートログを日付範囲を指定して削除できます。")
page_break(doc)

# =====================================================================
# 第16章
# =====================================================================
add_heading(doc, "第16章　スケジューラー・自動処理", 1)
add_para(doc, "システム内部でバックグラウンドに定期実行されるタスクの説明です。")

# 16.1: heading → table
add_heading(doc, "16.1　定期実行タスク一覧", 2)
add_table(doc,
    ["タスク名", "実行タイミング", "処理内容"],
    [
        ["発注点自動更新", "毎月1日", "reorder_auto=1：AI予測で更新\nreorder_auto=2：前年実績で更新"],
        ["売上日次集計", "毎日", "sales_history → sales_daily_agg へ集計（予測高速化）"],
        ["予測キャッシュ再構築", "随時", "forecast_cache にABC・予測値を事前計算してキャッシュ"],
        ["気象データ自動取得", "毎日03:00", "気象APIから観測地点の気温データを自動取得・保存"],
        ["CSV自動インポート", "設定した時刻", "run_times で指定した時刻に売上CSVを自動取込"],
        ["期限アラート確認", "毎日", "期限切れ間近の在庫を検出してアラートメール送信"],
    ],
    col_widths=[4, 3, 9]
)

add_heading(doc, "16.2　自動チェックの詳細動作", 2)
add_bullet(doc, "① 在庫が発注点以下の全有効商品を抽出")
add_bullet(doc, "② 混載グループなし → order_pending に追加（または即時発注）")
add_bullet(doc, "③ 混載グループあり → 同グループでまとめて条件判定")
add_bullet(doc, "④ 混載条件（ケース数）を満たした場合 → 発注処理")
add_bullet(doc, "⑤ 強制発注日数を超えたもの → 条件に関わらず強制発注")
add_bullet(doc, "⑥ 確定後 → order_history に記録・メール送信・ordered_at フラグをセット")
page_break(doc)

# =====================================================================
# 第17章
# =====================================================================
add_heading(doc, "第17章　権限一覧", 1)

# 17.1: heading → table
add_heading(doc, "17.1　ロールについて", 2)
add_table(doc,
    ["ロール", "説明"],
    [
        ["admin（管理者）", "全機能にアクセス可能。ユーザー管理・設定変更・データ削除も実施できる"],
        ["user（一般ユーザー）", "付与された権限の範囲でのみ操作可能。権限の変更はadminのみ実施可"],
    ],
    col_widths=[4, 12]
)

# 17.2: heading → para(kwn) → table
add_heading(doc, "17.2　権限キーと対応機能", 2)
add_para(doc, "各ユーザーに付与する権限キーと、アクセスできる機能の対応表です。", kwn=True)
add_table(doc,
    ["権限キー", "対応する主な操作"],
    [
        ["dashboard", "トップ画面（在庫サマリー）の閲覧"],
        ["products", "商品一覧の閲覧・検索（登録・編集・削除は admin のみ）"],
        ["orders", "発注操作（発注する・自動チェック・発注残管理・混載管理）"],
        ["inventory", "在庫一覧・直接編集・廃棄・ロケーション移動・ピッキング・棚補充"],
        ["forecast", "需要予測閲覧・発注点自動適用・販促計画・受注予定・ABC分析・気温・MD計画"],
        ["chains", "チェーン・店舗・仕入先CD・商品CD別設定の登録・編集"],
        ["order_history", "発注履歴閲覧・受領管理（受領入力・履歴）"],
        ["reports", "各種レポート閲覧"],
        ["receipt", "受領入力（発注残からの部分受領も含む）"],
        ["csv", "CSV取込設定の作成・実行・ログ確認"],
        ["stocktake", "棚卸実施・確定・エクスポート"],
        ["users", "ユーザー管理（admin相当）"],
        ["settings", "システム設定の変更・バックアップ・復元"],
    ],
    col_widths=[4, 12]
)
page_break(doc)

# =====================================================================
# 第18章
# =====================================================================
add_heading(doc, "第18章　用語集", 1)
add_table(doc,
    ["用語", "説明"],
    [
        ["JAN", "Japan Article Number。商品バーコードに使われる13桁の番号"],
        ["発注点（reorder_point）", "在庫がこの数量以下になると発注検討対象になるしきい値"],
        ["発注数量（reorder_qty）", "1回の発注で仕入先に送る数量"],
        ["入数（unit_qty）", "1ケース（段ボール）に入る個数"],
        ["リードタイム", "発注してから入荷するまでの日数"],
        ["安全係数（safety_factor）", "需要の変動幅に対して余裕を持たせる係数"],
        ["動的安全在庫（DSS）", "Dynamic Safety Stock。需要の統計的ばらつきから算出した安全在庫量"],
        ["ABC分析", "売上金額の構成比で商品をA/B/Cに分類する管理手法"],
        ["混載グループ", "同一の輸送ロットにまとめる商品のグループ識別子"],
        ["混載ロットモード", "gte: 合計が最低ケース数以上で発注 / unit: 倍数単位で発注"],
        ["reorder_auto", "0=手動 / 1=AIモード / 2=前年実績。発注点の自動更新方法"],
        ["前年実績モード（ly）", "前年同月の販売実績をもとに発注点を自動計算するモード"],
        ["AIモード（ai）", "統計・機械学習アルゴリズムで需要予測して発注点を自動計算するモード"],
        ["SMA", "Simple Moving Average（単純移動平均）。最も基本的な予測アルゴリズム"],
        ["ARIMA", "時系列予測モデル。季節性・トレンドを考慮した高度な予測"],
        ["ETS", "Exponential Smoothing。指数平滑法による予測モデル"],
        ["uplift_factor", "販促時の売上上昇係数（例: 1.5 = 通常の150%の売上を予測）"],
        ["ロケーションコード", "倉庫内の保管場所を識別するコード（例: A-01-03）"],
        ["在庫引き当て除外", "特定チェーン/店舗への出荷分を在庫計算から除外する設定"],
        ["order_pending", "混載条件判定中で発注保留中の状態"],
        ["stock_movements", "在庫の入出庫・調整・廃棄・移動を記録する監査ログ"],
        ["NSSM", "Non-Sucking Service Manager。FlaskアプリをWindowsサービスとして管理するツール"],
        ["forecast_cache", "予測計算結果をキャッシュして画面表示を高速化するテーブル"],
        ["sales_daily_agg", "sales_historyの日次集計テーブル。予測クエリを高速化するために使用"],
    ],
    col_widths=[5, 11]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("― 以上 ―")
set_font(r, size=11)

out = r"C:\Users\sato-mzk-002\inventory_system\在庫管理システム_取扱説明書.docx"
doc.save(out)
print(f"生成完了: {out}")
